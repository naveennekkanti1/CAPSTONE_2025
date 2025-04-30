from flask import Flask, render_template, request, redirect, url_for, session, flash, url_for,send_file,jsonify,Response
from flask_pymongo import PyMongo
from werkzeug.security import generate_password_hash, check_password_hash
from bson import ObjectId,errors
from werkzeug.utils import secure_filename
from datetime import datetime,timedelta
from pymongo import MongoClient,DESCENDING
import os,uuid,smtplib,threading,gridfs,logging,io,random,docx,qrcode,base64
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from flask_mail import Mail, Message
from google.oauth2.service_account import Credentials as ServiceAccountCredentials



app = Flask(__name__)
app.secret_key = 'your_secret_key_here'
app.config['MONGO_URI'] = app.config['MONGO_URI'] = "mongodb+srv://durganaveen:nekkanti@cluster0.8nibi9x.mongodb.net/RAPACT?retryWrites=true&w=majority&appName=Cluster0"
client = MongoClient(app.config['MONGO_URI'])
db = client['RAPACT']
users_collection = db['users']
appointments_collection = db['appointments']
meetings_collection=db['meetings']
enquiry_collection=db['enquiry_details']
fs = gridfs.GridFS(db)
SCOPES = ["https://www.googleapis.com/auth/calendar"]

#Flask-Mail Sending
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = 'srmcorporationservices@gmail.com'
app.config['MAIL_PASSWORD'] = 'bxxo qcvd njfj kcsa'
app.config['MAIL_DEFAULT_SENDER'] = 'srmcorporationservices@gmail.com'
otp_store = {}  
mail = Mail(app)

# File upload settings
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'jpg', 'jpeg', 'png'}

mongo = PyMongo(app)

# Ensure upload folder exists

# File type check function
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

# Home route
@app.route('/')
def home():
    user_logged_in = 'user_id' in session
    return render_template('index.html', user_logged_in=user_logged_in)

@app.route('/services')
def services():
    return render_template('services.html')

# About route
@app.route('/about')
def about():
    return render_template('about.html')

@app.route('/predictions')
def predictions():
    return render_template('report_prediction.html')


@app.route('/submit_enquiry', methods=['POST'])
def submit_enquiry():
    name = request.form.get('name')
    email = request.form.get('email')
    message = request.form.get('message')

    if name and email and message:
        # Store in MongoDB
        mongo.db.enquiry_details.insert_one({
            "name": name,
            "email": email,
            "message": message,
            "status": "pending",
            "date": datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
        })

        # Plain text fallback
        text_body = f"""
        Hello {name},

        Thank you for reaching out! We have received your enquiry and will get back to you soon.

        Your Message:
        {message}

        Best Regards,
        RapiACT! Team❤️
        """

        # HTML formatted email
        html_body = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <title>Enquiry Received</title>
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    line-height: 1.6;
                    color: #333333;
                    margin: 0;
                    padding: 0;
                }}
                .container {{
                    max-width: 600px;
                    margin: 0 auto;
                    padding: 20px;
                }}
                .header {{
                    background-color: #2196F3;
                    color: white;
                    padding: 20px;
                    text-align: center;
                    border-top-left-radius: 5px;
                    border-top-right-radius: 5px;
                }}
                .content {{
                    background-color: #f9f9f9;
                    padding: 20px;
                    border: 1px solid #dddddd;
                    border-bottom-left-radius: 5px;
                    border-bottom-right-radius: 5px;
                }}
                .section {{
                    background-color: #ffffff;
                    border: 1px solid #dddddd;
                    padding: 15px;
                    margin: 20px 0;
                    border-radius: 5px;
                }}
                .section-title {{
                    font-weight: bold;
                    margin-bottom: 5px;
                    color: #2196F3;
                }}
                .footer {{
                    margin-top: 20px;
                    text-align: center;
                    font-size: 12px;
                    color: #777777;
                }}
                .logo {{
                    font-size: 24px;
                    font-weight: bold;
                    color: white;
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    <div class="logo">RapiACT!</div>
                </div>
                <div class="content">
                    <p>Hello {name},</p>
                    <p>Thank you for reaching out! We have received your enquiry and will get back to you soon.</p>
                    
                    <div class="section">
                        <div class="section-title">Your Message:</div>
                        <div>{message}</div>
                    </div>
                    
                    <p>We appreciate your interest. Stay tuned for updates.</p>
                    <p>Best Regards,<br>The RapiACT! Team ❤️</p>
                </div>
                <div class="footer">
                    <p>This is an automated message, please do not reply to this email.</p>
                    <p>&copy; 2025 RapiACT! All rights reserved.</p>
                </div>
            </div>
        </body>
        </html>
        """

        # Send confirmation email
        msg = Message("Enquiry Received", recipients=[email])
        msg.body = text_body
        msg.html = html_body
        mail.send(msg)

        return redirect('/')  # Redirect to home page after submission

    return "Failed to submit enquiry", 400


@app.route("/enquiry_details")
def get_enquiry_details():
    """ Fetch all enquiries from MongoDB sorted by newest first. """
    enquiries = list(mongo.db.enquiry_details.find().sort("date", -1))

    for enquiry in enquiries:
        if 'date' in enquiry and isinstance(enquiry['date'], str):
            try:
                enquiry['date'] = datetime.strptime(enquiry['date'], "%Y-%m-%d %H:%M:%S")
            except ValueError:
                enquiry['date'] = None

    return render_template("enquiry_details.html", enquiries=enquiries)

@app.route("/mark_done", methods=["POST"])
def mark_done():
    """ Mark an enquiry as done and send a response email with styled inline HTML. """
    data = request.json
    enquiry_id = data.get("enquiry_id")
    response_text = data.get("response")

    if not enquiry_id or not response_text:
        return jsonify({"success": False, "message": "Invalid data"}), 400

    try:
        enquiry = mongo.db.enquiry_details.find_one({"_id": ObjectId(enquiry_id)})

        if not enquiry:
            return jsonify({"success": False, "message": "Enquiry not found"}), 404

        result = mongo.db.enquiry_details.update_one(
            {"_id": ObjectId(enquiry_id)},
            {"$set": {
                "status": "done",
                "response": response_text,
                "response_date": datetime.utcnow()
            }}
        )

        if result.modified_count == 1:
            user_email = enquiry.get("email")
            user_name = enquiry.get("name")

            if user_email:
                html_body = f"""
                <html>
                <body style="font-family: Arial, sans-serif; color: #333; margin: 0; padding: 0;">
                    <div style="max-width: 600px; margin: auto;">
                        <div style="background-color: #2196F3; color: white; padding: 20px; text-align: center; border-top-left-radius: 5px; border-top-right-radius: 5px;">
                            <h2 style="margin: 0; font-size: 24px;">RapiACT!</h2>
                        </div>
                        <div style="background-color: #f9f9f9; padding: 20px; border: 1px solid #ddd; border-bottom-left-radius: 5px; border-bottom-right-radius: 5px;">
                            <p>Hello {user_name},</p>
                            <p>Your enquiry has been marked as <strong>resolved</strong>. Below are the details:</p>

                            <div style="background-color: #fff; border: 1px solid #ddd; padding: 15px; margin: 20px 0; border-radius: 5px;">
                                <div style="font-weight: bold; margin-bottom: 5px; color: #2196F3;">Your Message:</div>
                                <div>{enquiry.get('message')}</div>
                            </div>

                            <div style="background-color: #fff; border: 1px solid #ddd; padding: 15px; margin: 20px 0; border-radius: 5px;">
                                <div style="font-weight: bold; margin-bottom: 5px; color: #2196F3;">Admin Response:</div>
                                <div>{response_text}</div>
                            </div>

                            <p>If you have any further questions, feel free to reach out.</p>
                            <p>Thank you,<br><strong>RapiACT! Team ❤️</strong></p>
                        </div>
                        <div style="text-align: center; font-size: 12px; color: #777; margin-top: 20px;">
                            <p>This is an automated message, please do not reply to this email.</p>
                            <p>&copy; 2025 RapiACT! All rights reserved.</p>
                        </div>
                    </div>
                </body>
                </html>
                """

                msg = Message(
                    subject="Your Enquiry has been Resolved",
                    recipients=[user_email]
                )
                msg.html = html_body
                mail.send(msg)

            return jsonify({"success": True, "message": "Enquiry marked as done and email sent."})

        return jsonify({"success": False, "message": "Failed to update enquiry."}), 500

    except Exception as e:
        return jsonify({"success": False, "message": str(e)}), 500



def send_email(subject, recipients, html_body):
    """Function to send emails with HTML only."""
    try:
        msg = Message(subject, recipients=recipients)
        msg.html = html_body  # Set only HTML content, no plain text
        mail.send(msg)
    except Exception as e:
        print(f"Error sending email: {e}")

def schedule_email(subject, recipients, body, send_time):
    """Schedule an email to be sent at a later time."""
    delay = (send_time - datetime.now()).total_seconds()
    if delay > 0:
        threading.Timer(delay, send_email, [subject, recipients, body]).start()
def get_google_credentials():
    creds = None
    token_path = 'token.json'  # Stores user's access token

    # Load existing credentials if available
    if os.path.exists(token_path):
        creds = Credentials.from_authorized_user_file(token_path, SCOPES)

    # If no valid credentials, prompt login
    if not creds or not creds.valid:
        flow = InstalledAppFlow.from_client_secrets_file(
            'templates/credentials.json', SCOPES)  # Use your OAuth 2.0 client secret
        creds = flow.run_local_server(port=0)

        # Save credentials for future use
        with open(token_path, 'w') as token:
            token.write(creds.to_json())

    return creds

def create_google_meet(patient_email, doctor_email, meeting_datetime):
    credentials = get_google_credentials()  # OAuth 2.0 authentication
    service = build('calendar', 'v3', credentials=credentials)

    event = {
        'summary': 'Doctor Consultation',
        'description': f'Meeting between {doctor_email} and {patient_email}',
        'start': {
            'dateTime': datetime.fromisoformat(meeting_datetime).isoformat(),
            'timeZone': 'Asia/Kolkata',
        },
        'end': {
            'dateTime': (datetime.fromisoformat(meeting_datetime) + timedelta(minutes=30)).isoformat(),
            'timeZone': 'Asia/Kolkata',
        },
        'attendees': [
            {'email': doctor_email},  # Doctor email
            {'email': patient_email}  # Patient email
        ],
        'conferenceData': {
            'createRequest': {
                'requestId': str(uuid.uuid4()),
                'conferenceSolutionKey': {
                    'type': 'hangoutsMeet'
                }
            }
        }
    }

    try:
        event = service.events().insert(
            calendarId='primary',
            body=event,
            conferenceDataVersion=1
        ).execute()

        return event.get('hangoutLink')  # Return Google Meet link
    except Exception as e:
        print(f"Error creating Google Meet: {e}")
        return None

@app.route('/schedule_meeting', methods=['GET', 'POST'])
def schedule_meeting():
    if 'user_id' in session and session['role'] in ['doctor', 'admin']:  
        doctor_id = ObjectId(session['user_id'])

        if request.method == 'POST':
            patient_id = ObjectId(request.form['patient_id'])
            meeting_datetime = request.form['meeting_datetime']

            # Validate if patient exists
            patient = mongo.db.users.find_one({'_id': patient_id}, {'name': 1, 'email': 1})
            if not patient or 'email' not in patient:
                flash("Selected patient does not exist or has no email.", "danger")
                return redirect(url_for('schedule_meeting'))

            patient_email = patient['email']

            # Validate if doctor exists
            doctor = mongo.db.users.find_one({'_id': doctor_id}, {'name': 1, 'email': 1})
            if not doctor or 'email' not in doctor:
                flash("Doctor account issue: No email found.", "danger")
                return redirect(url_for('schedule_meeting'))

            doctor_email = doctor['email']

            # Generate Google Meet link
            try:
                meeting_link = create_google_meet(patient_email, doctor_email, meeting_datetime)
            except Exception as e:
                flash(f"Error creating Google Meet: {str(e)}", "danger")
                return redirect(url_for('schedule_meeting'))

            # Generate a unique appointment ID
            appointment_id = str(ObjectId())

            # Store the meeting details with appointment ID
            meeting_id = meeting_link.split('/')[-1]  # Extract Google Meet ID
            mongo.db.meetings.insert_one({
                'appointment_id': appointment_id,  # Store unique appointment ID
                'doctor_id': doctor_id,
                'patient_id': patient_id,
                'meeting_id': meeting_id,
                'meeting_datetime': meeting_datetime,
                'meeting_link': meeting_link
            })

            # Send email notification
            # Send email notification with HTML and CSS
            subject = "📅 Your Medical Appointment is Scheduled"

                # HTML email body with CSS styling
            html_body = f"""
                <!DOCTYPE html>
<html>
<head>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 600px;
            margin: 0 auto;
        }}
        .email-container {{
            border: 1px solid #ddd;
            border-radius: 8px;
            padding: 20px;
            background-color: #f9f9f9;
        }}
        .header {{
            background-color: #4285F4;
            color: white;
            padding: 15px;
            text-align: center;
            border-radius: 5px 5px 0 0;
            margin-bottom: 20px;
        }}
        .appointment-details {{
            background-color: white;
            padding: 15px;
            border-radius: 5px;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
            margin-bottom: 20px;
        }}
        .detail-item {{
            margin-bottom: 10px;
            padding: 8px;
            border-bottom: 1px solid #eee;
        }}
        .detail-label {{
            font-weight: bold;
            color: #555;
        }}
        .meeting-button {{
            display: inline-block;
            background-color: #4CAF50;
            color: white;
            padding: 12px 20px;
            text-decoration: none;
            border-radius: 5px;
            font-weight: bold;
            margin: 10px 0;
        }}
        .footer {{
            text-align: center;
            margin-top: 20px;
            font-size: 14px;
            color: #777;
        }}
        .appointment-id {{
            background-color: #f1f1f1;
            padding: 8px;
            border-radius: 4px;
            font-family: monospace;
            font-size: 14px;
        }}
    </style>
</head>
<body>
    <div class="email-container">
        <div class="header">
            <h2>Medical Appointment Confirmation</h2>
        </div>
        
        <p>Dear {patient['name']},</p>
        
        <p>Your appointment with Dr. {doctor['name']} has been successfully scheduled.</p>
        
        <div class="appointment-details">
            <div class="detail-item">
                <span class="detail-label">📅 Date & Time:</span> {meeting_datetime}
            </div>
            <div class="detail-item">
                <span class="detail-label">👨‍⚕️ Doctor:</span> Dr. {doctor['name']}
            </div>
            <div class="detail-item">
                <span class="detail-label">🆔 Appointment ID:</span> 
                <span class="appointment-id">{appointment_id}</span>
            </div>
        </div>
        
        <div style="text-align: center;">
            <a href="{meeting_link}" class="meeting-button">Join Video Meeting</a>
        </div>
        
        <p>Please ensure you join on time and have a stable internet connection for the best experience.</p>
        
        <p>If you need to reschedule or have any questions, please contact us as soon as possible.</p>
        
        <div class="footer">
            <p>Best regards,<br>RapiACT! Team ❤️</p>
        </div>
    </div>
</body>
</html>
"""

# Send email with HTML content only
            send_email(subject, [patient_email], html_body)

            flash("Meeting scheduled successfully!", "success")
            return redirect(url_for('schedule_meeting'))

        # Fetch all patients for dropdown selection
        patients = list(mongo.db.users.find({'role': 'patient'}, {'_id': 1, 'name': 1}))

        # Fetch all scheduled meetings for the logged-in doctor
        meetings = list(mongo.db.meetings.find({'doctor_id': doctor_id}))
        doctor_name = mongo.db.users.find_one({'_id': doctor_id}, {'name': 1})

        # Separate meetings into ongoing, upcoming, and past
        ongoing_meetings = []
        upcoming_meetings = []
        past_meetings = []
        current_time = datetime.utcnow()

        for meeting in meetings:
            patient = mongo.db.users.find_one({'_id': meeting['patient_id']}, {'name': 1})
            meeting['patient_name'] = f"{patient['name']}" if patient else "Unknown"

            meeting_time = datetime.strptime(meeting['meeting_datetime'], "%Y-%m-%dT%H:%M")

            # Check if meeting is ongoing (within a 1-hour window)
            if meeting_time <= current_time <= meeting_time.replace(hour=meeting_time.hour + 1):
                ongoing_meetings.append(meeting)
            elif meeting_time > current_time:
                upcoming_meetings.append(meeting)
            else:
                past_meetings.append(meeting)

        return render_template('schedule_meeting.html', 
                               name=doctor_name.get('name', ''),
                               patients=patients, 
                               ongoing_meetings=ongoing_meetings, 
                               upcoming_meetings=upcoming_meetings, 
                               past_meetings=past_meetings)

    flash("Unauthorized access.", "danger")
    return redirect(url_for('login'))

def get_scheduled_meetings():
    meetings = db.meetings.find().sort("meeting_datetime", DESCENDING)
    
    meeting_list = []
    for meeting in meetings:
        meeting_list.append({
            "patient_name": f"{meeting.get('patient_name', '')}",
            "meeting_datetime": meetings.get("meeting_datetime", "Not Scheduled"),
            "meeting_link": meetings.get("meeting_link", "#")  # Default to "#" if link is missing
        })
    
    return meeting_list

@app.route('/join_meeting/<meeting_id>')
def join_meeting(meeting_id):
    if 'user_id' in session:
        meeting = mongo.db.meetings.find_one({"meeting_id": meeting_id})
        if meeting:
            return redirect(meeting['meeting_link'])  # Redirect to Google Meet
        flash("Meeting not found.", "danger")
        return redirect(url_for('patient_dashboard'))
    flash("Unauthorized access.", "danger")
    return redirect(url_for('login'))

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'GET':
        return render_template('register.html')

    elif request.method == 'POST':
        data = request.json
        email = data.get('email')
        password = data.get('password')
        confirm_password = data.get('confirm_password')
        role = data.get('role')

        if not email or not password or not confirm_password or not role:
            flash("All fields are required!", "danger")
            return redirect(url_for('register'))

        if password != confirm_password:
            flash("Passwords do not match!", "danger")
            return redirect(url_for('register'))

        if users_collection.find_one({"email": email}):
            flash("User already registered!", "danger")
            return redirect(url_for('register'))

        if email not in otp_store:
            flash("OTP not verified!", "danger")
            return redirect(url_for('register'))

        hashed_password = generate_password_hash(password)

        user_data = {
            "email": email,
            "password": hashed_password,
            "role": role,
            "created_at": datetime.utcnow()
        }
        users_collection.insert_one(user_data)
        del otp_store[email]

        flash("Registration successful!", "success")
        return redirect(url_for('login'))

@app.route('/send-otp', methods=['POST'])
def send_otp():
    data = request.json
    email = data.get('email')

    if not email:
        return jsonify({"error": "Email is required"}), 400
    
    if users_collection.find_one({"email": email}):
        return jsonify({"error": "User already registered!"}), 400

    otp = str(random.randint(100000, 999999))
    otp_store[email] = otp

    subject = "Your OTP Code for RapiACT Registration"
    
    # HTML email with CSS styling
    html_body = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            body {{
                font-family: Arial, sans-serif;
                line-height: 1.6;
                color: #333;
                max-width: 600px;
                margin: 0 auto;
            }}
            .email-container {{
                border: 1px solid #ddd;
                border-radius: 8px;
                padding: 20px;
                background-color: #f9f9f9;
            }}
            .header {{
                background-color: #6B46C1;
                color: white;
                padding: 15px;
                text-align: center;
                border-radius: 5px 5px 0 0;
                margin-bottom: 20px;
            }}
            .otp-container {{
                background-color: white;
                padding: 20px;
                border-radius: 5px;
                box-shadow: 0 2px 5px rgba(0,0,0,0.1);
                margin: 25px 0;
                text-align: center;
            }}
            .otp-code {{
                font-size: 32px;
                font-weight: bold;
                letter-spacing: 5px;
                color: #6B46C1;
                padding: 10px 20px;
                background-color: #f3f4f6;
                border-radius: 5px;
                margin: 10px 0;
                display: inline-block;
            }}
            .footer {{
                text-align: center;
                margin-top: 20px;
                font-size: 14px;
                color: #777;
            }}
            .expiry-notice {{
                background-color: #fffbeb;
                border-left: 4px solid #f59e0b;
                padding: 10px 15px;
                margin: 15px 0;
                font-size: 14px;
            }}
            .help-text {{
                font-size: 15px;
                color: #555;
            }}
        </style>
    </head>
    <body>
        <div class="email-container">
            <div class="header">
                <h2>One-Time Password (OTP)</h2>
            </div>
            
            <p>Hello,</p>
            
            <p>Thank you for registering with RapiACT! To verify your email address, please use the following One-Time Password (OTP):</p>
            
            <div class="otp-container">
                <div class="otp-code">{otp}</div>
                <p class="help-text">Please enter this code on the registration page.</p>
            </div>
            
            <div class="expiry-notice">
                <strong>Note:</strong> This OTP will expire in 10 minutes for security reasons.
            </div>
            
            <p>If you did not request this code, please ignore this email. Someone might have entered your email address by mistake.</p>
            
            <div class="footer">
                <p>Best regards,<br>RapiACT! Team ❤️</p>
                <p style="font-size: 12px; color: #999;">This is an automated message, please do not reply to this email.</p>
            </div>
        </div>
    </body>
    </html>
    """
    
    try:
        msg = Message(subject=subject, recipients=[email])
        msg.html = html_body
        mail.send(msg)
        return jsonify({"message": "OTP sent successfully!"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/verify-otp', methods=['POST'])
def verify_otp():
    data = request.json
    email = data.get('email')
    otp_entered = data.get('otp')
    name = data.get('name')
    password = data.get('password')
    role = data.get('role')

    if not email or not otp_entered:
        return jsonify({"error": "Email and OTP are required"}), 400

    if users_collection.find_one({"email": email}):
        return jsonify({"error": "User already registered!"}), 400  

    if email in otp_store and otp_store[email] == otp_entered:
        hashed_password = generate_password_hash(password)
        session.permanent = True  # Ensure session persists
        session["pending_user"] = {
            "name": name,
            "email": email,
            "password": hashed_password,
            "role": role
        }
        del otp_store[email]  # Remove OTP after successful verification

        return jsonify({"message": "OTP verified successfully!", "redirect_url": url_for('register')})

    return jsonify({"error": "Invalid OTP. Please try again!"}), 400

@app.route('/patient_register', methods=['GET', 'POST'])
def patient_register():
    if request.method == 'GET':
        return render_template('patient_register.html')

    elif request.method == 'POST':
        if "pending_user" not in session:
            return jsonify({"error": "No user session found. Please register again."}), 400

        data = request.form
        phone = data.get("phone")
        age = data.get("age")
        gender = data.get("gender")
        address = data.get("address")

        if not (phone and age and gender and address):
            return jsonify({"error": "All fields are required"}), 400

        # Get pending user details from session
        user_data = session["pending_user"]
        name = user_data.get("name")
        email = user_data.get("email")

        # Handle photo upload
        if 'photo' not in request.files:
            return jsonify({"error": "Photo is required"}), 400

        photo = request.files['photo']
        if photo.filename == '':
            return jsonify({"error": "Invalid photo file"}), 400

        # Store photo in GridFS
        photo_id = fs.put(photo, filename=photo.filename)

        # Complete user data
        user_data.update({
            "role": "patient",
            "phone": phone,
            "age": age,
            "gender": gender,
            "address": address,
            "photo": photo_id,  # Store reference to photo in MongoDB
            "created_at": datetime.utcnow()
        })

        # Save user in MongoDB
        users_collection.insert_one(user_data)

        # Remove session data
        session.pop("pending_user", None)

        # Send welcome email
        send_welcome_email1(name, email, "patient")

        return jsonify({"message": "Registration successful!", "redirect_url": url_for('login')})

@app.route('/doctor_register', methods=['GET', 'POST'])
def doctor_register():
    if request.method == 'GET':
        return render_template('doctor_register.html')

    elif request.method == 'POST':
        if "pending_user" not in session:
            return jsonify({"error": "No user session found. Please register again."}), 400

        data = request.form
        phone = data.get("phone")
        age = data.get("age")
        gender = data.get("gender")
        address = data.get("address")
        specialization = data.get("specialization")
        years_experience = data.get("experience")
        professional_experience = data.get("pro_experience", "")

        if not (phone and age and gender and specialization and years_experience and address):
            return jsonify({"error": "All fields are required."}), 400

        if 'photo' not in request.files:
            return jsonify({"error": "Photo is required"}), 400

        photo = request.files['photo']
        if photo.filename == '':
            return jsonify({"error": "Invalid photo file"}), 400

        photo_id = fs.put(photo, filename=photo.filename)

        # Retrieve email from session
        pending_user = session["pending_user"]
        name = pending_user.get("name")
        email = pending_user.get("email")
        password = pending_user.get("password")

        if not email:
            return jsonify({"error": "Email is missing. Please register again."}), 400

        user_data = {
            "email": email,  # Ensure email is saved
            "password": password,  # Save hashed password
            "name": name,
            "role": "doctor",
            "phone": phone,
            "age": age,
            "gender": gender,
            "address": address,
            "specialization": specialization,
            "years_experience": years_experience,
            "professional_experience": professional_experience,
            "photo": photo_id,
            "created_at": datetime.utcnow(),
            "account_status": "pending"
        }

        users_collection.insert_one(user_data)

        # Remove session data after registration
        session.pop("pending_user", None)

        # Send welcome email
        send_welcome_email1(name, email, "doctor")

        return jsonify({"message": "Registration successful! Await admin approval.", "redirect_url": url_for('login')})

def send_welcome_email1(name, email, role):
    try:
        # Configure email settings
        msg = Message(
            subject="Welcome to the RapiACT! Family ❤️",
            sender=app.config['MAIL_DEFAULT_SENDER'],
            recipients=[email]
        )
        
        # Customize message based on role
        role_specific_message = "book appointments and manage your healthcare needs" if role == "patient" else "manage patients and handle appointments"
        approval_message = "" if role == "patient" else "<p>Your doctor account is pending approval. An administrator will review your information and approve your account soon.</p>"
        
        # HTML email with styling and logo
        msg.html = f"""
<!DOCTYPE html>
<html>
<head>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            color: #333333;
            max-width: 600px;
            margin: 0 auto;
        }}
        .container {{
            padding: 20px;
            background-color: #f9f9f9;
            border-radius: 8px;
        }}
        .header {{
            background-color: #4a90e2;
            color: white;
            padding: 15px;
            text-align: center;
            border-radius: 8px 8px 0 0;
        }}
        .logo {{
            max-width: 150px;
            margin: 0 auto;
            display: block;
        }}
        .content {{
            background-color: white;
            padding: 20px;
            border-radius: 0 0 8px 8px;
            border: 1px solid #e0e0e0;
        }}
        .footer {{
            text-align: center;
            margin-top: 20px;
            color: #777;
            font-size: 14px;
        }}
        .highlight {{
            font-weight: bold;
            color: #4a90e2;
        }}
        .button {{
            display: inline-block;
            background-color: #4CAF50;
            color: white;
            padding: 12px 20px;
            text-decoration: none;
            border-radius: 4px;
            margin: 15px 0;
            font-weight: bold;
            text-align: center;
        }}
        .features {{
            margin: 20px 0;
        }}
        .feature-item {{
            margin-bottom: 10px;
            padding-left: 20px;
            position: relative;
        }}
        .feature-item:before {{
            content: "✓";
            color: #4CAF50;
            position: absolute;
            left: 0;
            font-weight: bold;
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <img src="https://capstone-2025-37jt.onrender.com/static/images/logo.jpg" alt="RapiACT! Logo" class="logo">
            <h2>Welcome to the RapiACT! Family</h2>
        </div>
        <div class="content">
            <h3>Hello, {name}!</h3>
            <p>We're thrilled to welcome you to the RapiACT! Healthcare System. Your account has been successfully created as a <span class="highlight">{role}</span>.</p>
            
            {approval_message}
            
            <p>As a {role}, you can now <span class="highlight">{role_specific_message}</span> through our intuitive platform.</p>
            
            <div style="text-align: center;">
                <a href="https://capstone-2025-37jt.onrender.com/login" class="button">Log In to Your Account</a>
            </div>
            
            <div class="features">
                <h4>What you can do with RapiACT!:</h4>
                <div class="feature-item">Access healthcare services quickly and efficiently</div>
                <div class="feature-item">Manage your medical records securely</div>
                <div class="feature-item">Schedule appointments with just a few clicks</div>
                <div class="feature-item">Receive timely notifications about your healthcare</div>
                <div class="feature-item">Connect with healthcare professionals effortlessly</div>
            </div>
            
            <p>We'll be sending you another email shortly with more details about our platform and a special welcome flyer!</p>
            
            <p>If you have any questions or need assistance, please don't hesitate to contact our support team at <a href="mailto:support@rapiact.com">support@rapiact.com</a>.</p>
        </div>
        <div class="footer">
            <p>Best regards,<br>The RapiACT! Team ❤️</p>
            <p>&copy; 2025 RapiACT! Healthcare. All rights reserved.</p>
        </div>
    </div>
</body>
</html>
"""
        # Send first welcome email
        mail.send(msg)
        
        # Send second email with details and flyer
        send_welcome_flyer_email(name, email, role)
        
        return True
    except Exception as e:
        print(f"Error sending welcome email: {str(e)}")
        return False

def send_welcome_flyer_email(name, email, role):
    try:
        # Configure email settings for the second email
        msg = Message(
            subject="Discover RapiACT! - Your Complete Healthcare Solution",
            sender=app.config['MAIL_DEFAULT_SENDER'],
            recipients=[email]
        )
        
        # HTML email with flyer and platform details
        msg.html = f"""
<!DOCTYPE html>
<html>
<head>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            color: #333333;
            max-width: 600px;
            margin: 0 auto;
        }}
        .container {{
            padding: 20px;
            background-color: #f9f9f9;
            border-radius: 8px;
        }}
        .header {{
            background-color: #4a90e2;
            color: white;
            padding: 15px;
            text-align: center;
            border-radius: 8px 8px 0 0;
        }}
        .logo {{
            max-width: 150px;
            margin: 0 auto;
            display: block;
        }}
        .content {{
            background-color: white;
            padding: 20px;
            border-radius: 0 0 8px 8px;
            border: 1px solid #e0e0e0;
        }}
        .flyer {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 20px auto;
            border: 1px solid #e0e0e0;
            border-radius: 8px;
        }}
        .section {{
            margin: 25px 0;
            padding-bottom: 15px;
            border-bottom: 1px solid #e0e0e0;
        }}
        .section-title {{
            color: #4a90e2;
            font-size: 18px;
            margin-bottom: 10px;
        }}
        .footer {{
            text-align: center;
            margin-top: 20px;
            color: #777;
            font-size: 14px;
        }}
        .social {{
            margin: 15px 0;
        }}
        .social a {{
            display: inline-block;
            margin: 0 10px;
            color: #4a90e2;
            text-decoration: none;
        }}
        .button {{
            display: inline-block;
            background-color: #4CAF50;
            color: white;
            padding: 12px 20px;
            text-decoration: none;
            border-radius: 4px;
            margin: 15px 0;
            font-weight: bold;
            text-align: center;
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <img src="https://capstone-2025-37jt.onrender.com/static/images/logo.jpg" alt="RapiACT! Logo" class="logo">
            <h2>Discover RapiACT! Healthcare</h2>
        </div>
        <div class="content">
            <h3>Dear {name},</h3>
            
            <p>Thank you for joining RapiACT! - where healthcare meets innovation. We're excited to have you as part of our growing community.</p>
            
            <div class="section">
                <h4 class="section-title">About RapiACT!</h4>
                <p>RapiACT! is a comprehensive healthcare platform designed to streamline the medical experience for both patients and healthcare providers. Our mission is to make quality healthcare accessible, efficient, and personalized for everyone.</p>
            </div>
            
            <div class="section">
                <h4 class="section-title">Key Features</h4>
                <ul>
                    <li><strong>Quick Appointments:</strong> Schedule appointments with just a few clicks</li>
                    <li><strong>Secure Communication:</strong> Chat directly with your healthcare provider</li>
                    <li><strong>Medical Records:</strong> Access and manage your medical history securely</li>
                    <li><strong>Reminders:</strong> Never miss an appointment with our notification system</li>
                    <li><strong>Health Tracking:</strong> Monitor your health metrics with our intuitive tools</li>
                </ul>
            </div>
            
            <div class="section">
                <h4 class="section-title">Getting Started</h4>
                <p>To make the most of RapiACT!, we recommend the following steps:</p>
                <ol>
                    <li>Complete your profile information</li>
                    <li>Explore the dashboard to familiarize yourself with available features</li>
                    <li>Download our mobile app for on-the-go access</li>
                    <li>Set your notification preferences</li>
                </ol>
                <div style="text-align: center;">
                    <a href="https://capstone-2025-37jt.onrender.com/guide" class="button">View User Guide</a>
                </div>
            </div>
            
            <p>We're constantly improving our services to better meet your healthcare needs. Your feedback is valuable to us!</p>
            
            <div class="social">
                <p>Connect with us:</p>
                <a href="https://facebook.com/rapiact">Facebook</a> |
                <a href="https://twitter.com/rapiact">Twitter</a> |
                <a href="https://instagram.com/rapiact">Instagram</a>
            </div>
        </div>
        <div class="footer">
            <p>Best regards,<br>The RapiACT! Team ❤️</p>
            <p>&copy; 2025 RapiACT! Healthcare. All rights reserved.</p>
            <p><small>If you have any questions or need assistance, contact us at support@rapiact.com</small></p>
        </div>
    </div>
</body>
</html>
"""
        # Send second email with flyer and details
        mail.send(msg)
        return True
    except Exception as e:
        print(f"Error sending flyer email: {str(e)}")
        return False

# ---- Login ----
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'GET':
        return render_template('login.html')

    elif request.method == 'POST':
        data = request.form
        email = data.get("email")
        password = data.get("password")
        captcha_input = data.get("captcha")
        captcha_value = data.get("captchaValue")
        remember_me = True if data.get("rememberMe") else False

        # Validate required fields
        if not (email and password):
            flash("Email and password are required", "error")
            return redirect(url_for('login'))
            
        # Validate CAPTCHA
        if not captcha_input or captcha_input != captcha_value:
            flash("CAPTCHA verification failed. Please try again.", "error")
            return redirect(url_for('login'))

        # Check if user exists and password is correct
        user = users_collection.find_one({"email": email})

        if not user or not check_password_hash(user["password"], password):
            flash("Invalid credentials", "error")
            return redirect(url_for('login'))

        # Check if doctor account is approved
        if user.get("role") == "doctor" and user.get("account_status") != "approved":
            flash("Your account is pending approval from the administrator.", "warning")
            return redirect(url_for('login'))

        # Set session variables
        session["user_id"] = str(user["_id"])
        session["role"] = user.get("role")
        
        # Set session permanence based on remember me checkbox
        session.permanent = remember_me

        flash("You have successfully logged in!", "success")
        return redirect(url_for('dashboard'))


# ---- Admin Dashboard ----
@app.route('/admin_dashboard')
def admin_dashboard():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash("Unauthorized access.", "danger")
        return redirect(url_for('login'))

    total_users = users_collection.count_documents({})
    total_doctors = users_collection.count_documents({'role': 'doctor'})
    total_patients = users_collection.count_documents({'role': 'patient'})
    total_appointments=appointments_collection.count_documents({})
    total_meetings=meetings_collection.count_documents({})
    total_enquires=enquiry_collection.count_documents({})

    unapproved_doctors = list(users_collection.find({'role': 'doctor', 'account_status': 'pending'},
                                                     {'_id': 1, 'name': 1, 'email': 1,'specialization':1,'years_experience':1,'professional_experience':1,'age':1,'gender':1}))

    return render_template('admin_dashboard.html',
                           total_users=total_users,
                           total_doctors=total_doctors,
                           total_patients=total_patients,
                           total_appointments=total_appointments,
                           total_meetings=total_meetings,
                           total_enquires=total_enquires,
                           unapproved_doctors=unapproved_doctors)

import secrets


@app.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'GET':
        return render_template('forgot_password.html', step="email")
    
    elif request.method == 'POST':
        email = request.form.get('email')
        
        if not email:
            flash("Email is required", "error")
            return redirect(url_for('forgot_password'))
        
        # Check if user exists
        user = users_collection.find_one({"email": email})
        if not user:
            # Still show success message for security reasons
            flash("If your email is registered, you will receive password reset instructions", "info")
            return redirect(url_for('login'))
        
        # Generate OTP
        otp = str(random.randint(100000, 999999))
        otp_store[email] = {
            'otp': otp,
            'expires_at': datetime.utcnow() + timedelta(minutes=10)
        }
        
        # HTML Email Template with CSS styling
        html_body = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Reset Your Password</title>
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    line-height: 1.6;
                    color: #333333;
                    margin: 0;
                    padding: 0;
                }}
                .container {{
                    max-width: 600px;
                    margin: 0 auto;
                    padding: 20px;
                }}
                .header {{
                    background-color: #4CAF50;
                    color: white;
                    padding: 20px;
                    text-align: center;
                    border-top-left-radius: 5px;
                    border-top-right-radius: 5px;
                }}
                .content {{
                    background-color: #f9f9f9;
                    padding: 20px;
                    border: 1px solid #dddddd;
                    border-bottom-left-radius: 5px;
                    border-bottom-right-radius: 5px;
                }}
                .otp-container {{
                    background-color: #ffffff;
                    border: 1px solid #dddddd;
                    padding: 15px;
                    margin: 20px 0;
                    text-align: center;
                    border-radius: 5px;
                }}
                .otp-code {{
                    font-size: 32px;
                    font-weight: bold;
                    letter-spacing: 5px;
                    color: #4CAF50;
                }}
                .footer {{
                    margin-top: 20px;
                    text-align: center;
                    font-size: 12px;
                    color: #777777;
                }}
                .expire-note {{
                    color: #FF5722;
                    font-size: 14px;
                    text-align: center;
                    margin-top: 10px;
                }}
                .logo {{
                    font-size: 24px;
                    font-weight: bold;
                    color: white;
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    <div class="logo">RapiACT!</div>
                </div>
                <div class="content">
                    <p>Hello,</p>
                    <p>We received a request to reset your password for your RapiACT! account. To proceed with the password reset, please use the following One-Time Password (OTP):</p>
                    
                    <div class="otp-container">
                        <div class="otp-code">{otp}</div>
                    </div>
                    
                    <p class="expire-note">This code will expire in 10 minutes.</p>
                    
                    <p>If you did not request a password reset, please ignore this email and your account will remain secure.</p>
                    
                    <p>Thank you,<br>The RapiACT! Team ❤️</p>
                </div>
                <div class="footer">
                    <p>This is an automated message, please do not reply to this email.</p>
                    <p>&copy; 2025 RapiACT! All rights reserved.</p>
                </div>
            </div>
        </body>
        </html>
        """
        
        # Plain text version as fallback
        text_body = f"""
        Hello,

        You have requested to reset your password for your RapiACT! account. Your OTP code is:

        {otp}

        This code will expire in 10 minutes.

        If you did not request a password reset, please ignore this email.

        Best Regards,
        RapiACT! Team❤️
        """
        
        try:
            msg = Message(subject="Reset Your Password - RapiACT!", recipients=[email])
            msg.body = text_body
            msg.html = html_body
            mail.send(msg)
            
            # Store email in session for next step
            session['reset_email'] = email
            
            flash("OTP sent to your email", "success")
            return render_template('forgot_password.html', step="otp")
        except Exception as e:
            flash(f"Error sending email: {str(e)}", "error")
            return redirect(url_for('forgot_password'))

@app.route('/verify-reset-otp', methods=['POST'])
def verify_reset_otp():
    if 'reset_email' not in session:
        flash("Please start the password reset process again", "error")
        return redirect(url_for('forgot_password'))
    
    email = session.get('reset_email')
    otp_entered = request.form.get('otp')
    
    if not otp_entered:
        flash("OTP is required", "error")
        return render_template('forgot_password.html', step="otp")
    
    # Verify OTP
    if email not in otp_store or otp_store[email]['otp'] != otp_entered:
        flash("Invalid OTP", "error")
        return render_template('forgot_password.html', step="otp")
    
    # Check if OTP is expired
    if otp_store[email]['expires_at'] < datetime.utcnow():
        flash("OTP has expired", "error")
        del otp_store[email]
        return redirect(url_for('forgot_password'))
    
    # Generate reset token
    reset_token = secrets.token_urlsafe(32)
    token_expiry = datetime.utcnow() + timedelta(hours=1)
    
    # Store token in database
    users_collection.update_one(
        {"email": email},
        {"$set": {
            "reset_token": reset_token,
            "reset_token_expiry": token_expiry
        }}
    )
    
    # Remove OTP after use
    del otp_store[email]
    
    # Set token in session
    session['reset_token'] = reset_token
    
    flash("OTP verified successfully", "success")
    return render_template('forgot_password.html', step="new-password")

@app.route('/reset-password', methods=['POST'])
def reset_password():
    if 'reset_email' not in session or 'reset_token' not in session:
        flash("Please start the password reset process again", "error")
        return redirect(url_for('forgot_password'))
    
    email = session.get('reset_email')
    token = session.get('reset_token')
    
    # Verify token is still valid in the database
    user = users_collection.find_one({
        "email": email,
        "reset_token": token,
        "reset_token_expiry": {"$gt": datetime.utcnow()}
    })
    
    if not user:
        flash("Password reset link has expired", "error")
        return redirect(url_for('forgot_password'))
    
    new_password = request.form.get('new_password')
    confirm_password = request.form.get('confirm_password')
    
    if not (new_password and confirm_password):
        flash("All fields are required", "error")
        return render_template('forgot_password.html', step="new-password")
    
    if new_password != confirm_password:
        flash("Passwords do not match", "error")
        return render_template('forgot_password.html', step="new-password")
    
    # Update password
    hashed_password = generate_password_hash(new_password)
    
    users_collection.update_one(
        {"email": email},
        {"$set": {
            "password": hashed_password,
            "reset_token": None,
            "reset_token_expiry": None
        }}
    )
    
    # Clear session data
    session.pop('reset_email', None)
    session.pop('reset_token', None)
    
    flash("Password has been reset successfully. You can now log in with your new password.", "success")
    return redirect(url_for('login'))

# ---- Approve Doctor ----
@app.route('/approve_doctor/<doctor_id>', methods=['POST'])
def approve_doctor(doctor_id):
    if 'user_id' not in session or session.get('role') != 'admin':
        flash("Unauthorized access", "danger")
        return redirect(url_for('admin_dashboard'))

    doctor = users_collection.find_one({'_id': ObjectId(doctor_id)}, {'email': 1, 'name':1})

    if not doctor:
        flash("Doctor not found", "danger")
        return redirect(url_for('admin_dashboard'))

    if 'email' not in doctor:
        flash("Doctor email missing in database", "danger")
        return redirect(url_for('admin_dashboard'))

    # Approve the doctor
    users_collection.update_one({'_id': ObjectId(doctor_id)}, {'$set': {'account_status': 'approved'}})

    subject = "Doctor Registration Approved"
    recipients = [doctor['email']]
    
    # HTML email with CSS styling
    html_body = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            body {{
                font-family: Arial, sans-serif;
                line-height: 1.6;
                color: #333;
                max-width: 600px;
                margin: 0 auto;
            }}
            .email-container {{
                border: 1px solid #ddd;
                border-radius: 8px;
                padding: 20px;
                background-color: #f9f9f9;
            }}
            .header {{
                background-color: #4285F4;
                color: white;
                padding: 15px;
                text-align: center;
                border-radius: 5px 5px 0 0;
                margin-bottom: 20px;
            }}
            .content {{
                background-color: white;
                padding: 20px;
                border-radius: 5px;
                box-shadow: 0 2px 5px rgba(0,0,0,0.1);
                margin-bottom: 20px;
            }}
            .footer {{
                text-align: center;
                margin-top: 20px;
                font-size: 14px;
                color: #777;
            }}
            .login-button {{
                display: inline-block;
                background-color: #4CAF50;
                color: white;
                padding: 12px 24px;
                text-decoration: none;
                border-radius: 5px;
                font-weight: bold;
                margin: 10px 0;
            }}
        </style>
    </head>
    <body>
        <div class="email-container">
            <div class="header">
                <h2>Registration Approved! 🎉</h2>
            </div>
            
            <div class="content">
                <p>Hello Dr. {doctor.get('name', 'Doctor')},</p>
                
                <p>We are pleased to inform you that your registration with RapiACT has been <strong>approved</strong>!</p>
                
                <p>You can now log in to your account and start using our platform to manage appointments and connect with patients.</p>
                
                <div style="text-align: center; margin: 25px 0;">
                    <a href="https://capstone-2025-37jt.onrender.com//login" class="login-button">Login to Your Account</a>
                </div>
                
                <p>If you have any questions or need assistance, please don't hesitate to contact our support team.</p>
            </div>
            
            <div class="footer">
                <p>Thank you for joining us!<br>RapiACT! Team ❤️</p>
            </div>
        </div>
    </body>
    </html>
    """

    send_email(subject, recipients, html_body)

    flash("Doctor approved successfully! An email notification has been sent.", "success")
    return redirect(url_for('admin_dashboard'))


@app.route('/reject_doctor/<doctor_id>', methods=['POST'])
def reject_doctor(doctor_id):
    if 'user_id' not in session or session.get('role') != 'admin':
        flash("Unauthorized access", "danger")
        return redirect(url_for('admin_dashboard'))

    doctor = users_collection.find_one({'_id': ObjectId(doctor_id)}, {'email': 1, 'name': 1})

    if not doctor:
        flash("Doctor not found", "danger")
        return redirect(url_for('admin_dashboard'))

    if 'email' not in doctor:
        flash("Doctor email missing in database", "danger")
        return redirect(url_for('admin_dashboard'))

    # Delete the doctor from the database
    users_collection.delete_one({'_id': ObjectId(doctor_id)})

    subject = "Doctor Registration Status"
    recipients = [doctor['email']]
    
    # HTML email with CSS styling
    html_body = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            body {{
                font-family: Arial, sans-serif;
                line-height: 1.6;
                color: #333;
                max-width: 600px;
                margin: 0 auto;
            }}
            .email-container {{
                border: 1px solid #ddd;
                border-radius: 8px;
                padding: 20px;
                background-color: #f9f9f9;
            }}
            .header {{
                background-color: #e74c3c;
                color: white;
                padding: 15px;
                text-align: center;
                border-radius: 5px 5px 0 0;
                margin-bottom: 20px;
            }}
            .content {{
                background-color: white;
                padding: 20px;
                border-radius: 5px;
                box-shadow: 0 2px 5px rgba(0,0,0,0.1);
                margin-bottom: 20px;
            }}
            .footer {{
                text-align: center;
                margin-top: 20px;
                font-size: 14px;
                color: #777;
            }}
            .contact-button {{
                display: inline-block;
                background-color: #3498db;
                color: white;
                padding: 12px 24px;
                text-decoration: none;
                border-radius: 5px;
                font-weight: bold;
                margin: 10px 0;
            }}
        </style>
    </head>
    <body>
        <div class="email-container">
            <div class="header">
                <h2>Registration Update</h2>
            </div>
            
            <div class="content">
                <p>Hello Dr. {doctor.get('name', 'Doctor')},</p>
                
                <p>Thank you for your interest in joining the RapiACT platform.</p>
                
                <p>After careful review, we regret to inform you that your registration request has not been approved at this time.</p>
                
                <p>This decision may be due to one of the following reasons:</p>
                <ul>
                    <li>Incomplete or inaccurate information provided</li>
                    <li>Unable to verify professional credentials</li>
                    <li>High volume of applications in your specialty area</li>
                </ul>
                
                <p>If you believe this decision was made in error or would like more information, please contact our support team.</p>
                
                <div style="text-align: center; margin: 25px 0;">
                    <a href="mailto:support@rapidact.com" class="contact-button">Contact Support</a>
                </div>
            </div>
            
            <div class="footer">
                <p>We wish you the best in your professional endeavors.<br>RapiACT! Team</p>
            </div>
        </div>
    </body>
    </html>
    """

    send_email(subject, recipients, html_body)

    flash("Doctor rejected and removed successfully. An email notification has been sent.", "info")
    return redirect(url_for('admin_dashboard'))

@app.route('/user_photo/<user_id>')
def user_photo(user_id):
    try:
        # Ensure user_id is a valid ObjectId
        user_id = ObjectId(user_id)
    except errors.InvalidId:
        return send_file("static/images/logo.jpg", mimetype="image/png")  # Return default image if invalid ID

    user = db.users.find_one({"_id": user_id})
    
    if user and "photo" in user:
        try:
            image = fs.get(ObjectId(user["photo"]))  # Fetch image from GridFS
            return send_file(io.BytesIO(image.read()), mimetype="image/jpeg")
        except Exception as e:
            print(f"Error fetching image: {e}")
    
    return send_file("static/images/logo.jpg", mimetype="image/png")  # Default image if error

# Dashboard route
@app.route('/dashboard')
def dashboard():
    if 'user_id' in session:
        role = session.get('role')
        if role == 'admin':
            return redirect(url_for('admin_dashboard'))
        elif role == 'doctor':
            return redirect(url_for('doctor_dashboard'))
        elif role == 'patient':
            return redirect(url_for('patient_dashboard'))
    flash("Please log in first.", "danger")
    return redirect(url_for('login'))
# admin dashboard

@app.route('/users')
def users():
    role = request.args.get('role')
    query = {"role": role} if role else {}

    users = list(mongo.db.users.find(query))
    
    # Convert ObjectId to string for frontend display
    for user in users:
        user['_id'] = str(user['_id']) 

    return render_template('users.html', users=users)

@app.route('/appointments')
def appointments():
    # Fetch appointments
    appointments = list(db.appointments.find())
    
    # Fetch doctor details and enrich appointments data
    for appointment in appointments:
        doctor = db.doctors.find_one({"_id": appointment["doctor_id"]})
        if doctor:
            appointment["doctor_name"] = doctor.get("name", "Unknown")
            appointment["doctor_specialization"] = doctor.get("specialization", "Unknown")
        else:
            appointment["doctor_name"] = "Unknown"
            appointment["doctor_specialization"] = "Unknown"
    
    return render_template('appointments.html', appointments=appointments)

@app.route('/patient_dashboard')
@app.route('/patient_dashboard/<appointment_type>')
def patient_dashboard(appointment_type=None):
    if 'user_id' not in session or session['role'] != 'patient':
        flash("Unauthorized access.", "danger")
        return redirect(url_for('login'))

    patient_id = ObjectId(session['user_id'])

    # Fetch patient details
    patient = mongo.db.users.find_one({'_id': patient_id}, {'name': 1})
    if not patient:
        flash("Patient record not found.", "danger")
        return redirect(url_for('login'))

    # Fetch patient's appointments and meetings
    appointments = list(mongo.db.appointments.find({"patient_id": patient_id}))
    meetings = list(mongo.db.meetings.find({"patient_id": patient_id}))

    # Fetch doctor details for each appointment
    for appointment in appointments:
        doctor = mongo.db.users.find_one({"_id": appointment["doctor_id"]}, {"name":1, "specialization": 1})
        if doctor:
            appointment["doctor_name"] = doctor.get("name")
            appointment["specialization"] = doctor.get("specialization")
        else:
            appointment["doctor_name"] = "Unknown"
            appointment["specialization"] = "Unknown"

    # Classify appointments
    upcoming_appointments, ongoing_appointments, completed_appointments = [], [], []
    current_time = datetime.utcnow()

    for appointment in appointments:
        try:
            appointment_time = datetime.strptime(appointment['appointment_datetime'], '%Y-%m-%dT%H:%M')
        except ValueError:
            flash("Invalid appointment date format.", "danger")
            continue

        if appointment.get("status") in ["done", "completed"]:
            completed_appointments.append(appointment)
        elif appointment_time > current_time:
            upcoming_appointments.append(appointment)
        else:
            ongoing_appointments.append(appointment)

    # Filter by `appointment_type`
    appointment_filters = {
        "upcoming": (upcoming_appointments, [], []),
        "ongoing": ([], ongoing_appointments, []),
        "completed": ([], [], completed_appointments)
    }

    filtered_appointments = appointment_filters.get(appointment_type, 
        (upcoming_appointments, ongoing_appointments, completed_appointments))

    return render_template(
    'patient_dashboard.html',
    name=patient.get('name', ''),
    user=patient,  # Pass the entire patient object instead of just 'name'
    upcoming_appointments=filtered_appointments[0],
    ongoing_appointments=filtered_appointments[1],
    completed_appointments=filtered_appointments[2],
    meetings=meetings
)


def get_doctor_info():
    if 'user_id' in session and session['role'] == 'doctor':
        doctor_id = ObjectId(session['user_id'])
        doctor = mongo.db.users.find_one({'_id': doctor_id}, {'name': 1})
        return doctor
    return None

@app.route('/doctor_dashboard')
@app.route('/doctor_dashboard/<appointment_type>')
def doctor_dashboard(appointment_type=None):
    if 'user_id' in session and session['role'] == 'doctor':
        doctor_id = ObjectId(session['user_id'])
        doctor = get_doctor_info()
        doctor_name = mongo.db.users.find_one({'_id': doctor_id}, {'name': 1})
        
        if not doctor:
            flash("Doctor record not found.", "danger")
            return redirect(url_for('login'))
        
        appointments = list(mongo.db.appointments.find({"doctor_id": doctor_id}))
        upcoming_appointments = []
        ongoing_appointments = []
        completed_appointments = []
        current_time = datetime.now()
        
        for appointment in appointments:
            appointment_time = datetime.strptime(appointment['appointment_datetime'], '%Y-%m-%dT%H:%M')
            
            if appointment.get("status") in ["done", "completed"]:
                completed_appointments.append(appointment)
            elif appointment_time > current_time:
                upcoming_appointments.append(appointment)
            else:
                ongoing_appointments.append(appointment)
        
        # Filter based on appointment_type
        if appointment_type == "upcoming":
            return render_template('doctor_dashboard.html', doctor=doctor,
                                   name=doctor_name.get('name', ''),
                                   upcoming_appointments=upcoming_appointments, 
                                   ongoing_appointments=[], 
                                   completed_appointments=[])
        elif appointment_type == "ongoing":
            return render_template('doctor_dashboard.html', doctor=doctor,
                                   name=doctor_name.get('name', ''),
                                   upcoming_appointments=[], 
                                   ongoing_appointments=ongoing_appointments, 
                                   completed_appointments=[])
        elif appointment_type == "completed":
            return render_template('doctor_dashboard.html', doctor=doctor,
                                   name=doctor_name.get('name', ''),
                                   upcoming_appointments=[], 
                                   ongoing_appointments=[], 
                                   completed_appointments=completed_appointments)
        
        # Default: Show all
        return render_template('doctor_dashboard.html', doctor=doctor,
                               name=doctor_name.get('name', ''),
                               upcoming_appointments=upcoming_appointments,
                               ongoing_appointments=ongoing_appointments,
                               completed_appointments=completed_appointments)
    return redirect(url_for('login'))
# Route to cancel an appointment
@app.route('/cancel_appointment/<appointment_id>', methods=['POST'])
def cancel_appointment(appointment_id):
    if 'user_id' in session:
        appointment = mongo.db.appointments.find_one({'_id': ObjectId(appointment_id)})
        
        if appointment and str(appointment['patient_id']) == session['user_id']:
            # Get doctor information
            doctor = mongo.db.users.find_one({'_id': appointment['doctor_id']})
            doctor_name = doctor['name'] if doctor else "your doctor"
            doctor_email = doctor['email'] if doctor else None
            
            # Format datetime for display
            appointment_datetime = appointment.get('appointment_datetime')
            formatted_datetime = ""
            if appointment_datetime:
                try:
                    formatted_datetime = datetime.strptime(
                        appointment_datetime, "%Y-%m-%dT%H:%M"
                    ).strftime("%A, %B %d, %Y at %I:%M %p")
                except ValueError:
                    formatted_datetime = appointment_datetime  # Use as is if format is different
            
            # Delete the appointment
            mongo.db.appointments.delete_one({'_id': ObjectId(appointment_id)})
            
            # Generate dashboard URL for email template
            dashboard_url = request.host_url.rstrip('/') + url_for('patient_dashboard')
            doctor_dashboard_url = request.host_url.rstrip('/') + url_for('doctor_dashboard')
            
            # HTML Email Template for Patient Cancellation
            patient_email_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>Appointment Cancellation</title>
                <style>
                    body {{
                        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                        line-height: 1.6;
                        color: #333;
                        margin: 0;
                        padding: 0;
                    }}
                    .email-container {{
                        max-width: 600px;
                        margin: 0 auto;
                        padding: 20px;
                    }}
                    .header {{
                        background-color: #ea4335;
                        color: white;
                        padding: 20px;
                        text-align: center;
                        border-radius: 5px 5px 0 0;
                    }}
                    .content {{
                        background-color: #ffffff;
                        padding: 20px;
                        border-left: 1px solid #e0e0e0;
                        border-right: 1px solid #e0e0e0;
                    }}
                    .cancellation-details {{
                        background-color: #fef8f8;
                        padding: 15px;
                        border-radius: 5px;
                        margin: 15px 0;
                        border-left: 4px solid #ea4335;
                    }}
                    .cancellation-item {{
                        margin-bottom: 10px;
                    }}
                    .cancellation-label {{
                        font-weight: bold;
                        color: #ea4335;
                    }}
                    .footer {{
                        background-color: #f5f5f5;
                        padding: 15px;
                        text-align: center;
                        font-size: 12px;
                        color: #757575;
                        border-radius: 0 0 5px 5px;
                        border: 1px solid #e0e0e0;
                    }}
                    .button {{
                        display: inline-block;
                        background-color: #4285f4;
                        color: white;
                        padding: 10px 20px;
                        text-decoration: none;
                        border-radius: 5px;
                        margin-top: 15px;
                    }}
                    .logo {{
                        font-size: 24px;
                        font-weight: bold;
                        color: white;
                    }}
                </style>
            </head>
            <body>
                <div class="email-container">
                    <div class="header">
                        <div class="logo">RapiACT!</div>
                        <h2>Appointment Cancellation</h2>
                    </div>
                    <div class="content">
                        <p>Dear {appointment['patient_name']},</p>
                        <p>Your appointment has been successfully canceled.</p>
                        
                        <div class="cancellation-details">
                            <div class="cancellation-item">
                                <span class="cancellation-label">Canceled Appointment:</span> {formatted_datetime}
                            </div>
                            <div class="cancellation-item">
                                <span class="cancellation-label">Doctor:</span> Dr. {doctor_name}
                            </div>
                            <div class="cancellation-item">
                                <span class="cancellation-label">Reason for Visit:</span> {appointment.get('cause', 'Not specified')}
                            </div>
                        </div>
                        
                        <p>If you would like to schedule a new appointment, please visit your dashboard.</p>
                        
                        <center>
                            <a href="{dashboard_url}" class="button">Schedule New Appointment</a>
                        </center>
                    </div>
                    <div class="footer">
                        <p>© 2025 RapiACT! Healthcare Services</p>
                        <p>This is an automated email. Please do not reply to this message.</p>
                    </div>
                </div>
            </body>
            </html>
            """
            
            # HTML Email Template for Doctor Cancellation Notification
            doctor_email_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>Appointment Cancellation Notice</title>
                <style>
                    body {{
                        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                        line-height: 1.6;
                        color: #333;
                        margin: 0;
                        padding: 0;
                    }}
                    .email-container {{
                        max-width: 600px;
                        margin: 0 auto;
                        padding: 20px;
                    }}
                    .header {{
                        background-color: #ea4335;
                        color: white;
                        padding: 20px;
                        text-align: center;
                        border-radius: 5px 5px 0 0;
                    }}
                    .content {{
                        background-color: #ffffff;
                        padding: 20px;
                        border-left: 1px solid #e0e0e0;
                        border-right: 1px solid #e0e0e0;
                    }}
                    .cancellation-details {{
                        background-color: #fef8f8;
                        padding: 15px;
                        border-radius: 5px;
                        margin: 15px 0;
                        border-left: 4px solid #ea4335;
                    }}
                    .cancellation-item {{
                        margin-bottom: 10px;
                    }}
                    .cancellation-label {{
                        font-weight: bold;
                        color: #ea4335;
                    }}
                    .footer {{
                        background-color: #f5f5f5;
                        padding: 15px;
                        text-align: center;
                        font-size: 12px;
                        color: #757575;
                        border-radius: 0 0 5px 5px;
                        border: 1px solid #e0e0e0;
                    }}
                    .button {{
                        display: inline-block;
                        background-color: #34a853;
                        color: white;
                        padding: 10px 20px;
                        text-decoration: none;
                        border-radius: 5px;
                        margin-top: 15px;
                    }}
                    .logo {{
                        font-size: 24px;
                        font-weight: bold;
                        color: white;
                    }}
                </style>
            </head>
            <body>
                <div class="email-container">
                    <div class="header">
                        <div class="logo">RapiACT!</div>
                        <h2>Appointment Cancellation Notice</h2>
                    </div>
                    <div class="content">
                        <p>Dear Dr. {doctor_name},</p>
                        <p>An appointment has been canceled.</p>
                        
                        <div class="cancellation-details">
                            <div class="cancellation-item">
                                <span class="cancellation-label">Patient:</span> {appointment['patient_name']}
                            </div>
                            <div class="cancellation-item">
                                <span class="cancellation-label">Canceled Appointment:</span> {formatted_datetime}
                            </div>
                            <div class="cancellation-item">
                                <span class="cancellation-label">Reason for Visit:</span> {appointment.get('cause', 'Not specified')}
                            </div>
                        </div>
                        
                        <p>Your schedule has been updated accordingly.</p>
                        
                        <center>
                            <a href="{doctor_dashboard_url}" class="button">View Updated Schedule</a>
                        </center>
                    </div>
                    <div class="footer">
                        <p>© 2025 RapiACT! Healthcare Services</p>
                        <p>This is an automated email. Please do not reply to this message.</p>
                    </div>
                </div>
            </body>
            </html>
            """

            # Send HTML emails
            email = appointment['email']
            subject = "Appointment Cancellation"
            send_email1(subject, [email], None, patient_email_html)
            
            # Also notify the doctor if email is available
            if doctor_email:
                send_email1("Appointment Cancellation Notice", [doctor_email], None, doctor_email_html)

            flash("Appointment canceled successfully.", "success")
        else:
            flash("Unauthorized action.", "danger")
        
        return redirect(url_for('patient_dashboard'))
    
    return redirect(url_for('login'))
from bson.binary import Binary
def get_content_type(file_ext):
    content_types = {
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.png': 'image/png',
        '.gif': 'image/gif'
    }
    return content_types.get(file_ext, 'image/jpeg')

@app.route('/profile', methods=['GET', 'POST'])
def profile():
    user_id = session.get('user_id')

    if not user_id:
        return redirect(url_for('login'))

    try:
        user = mongo.db.users.find_one({"_id": ObjectId(user_id)})

        if not user:
            return "User not found", 404

        if request.method == 'POST':
            if 'update_profile' in request.form:
                # Handle profile update
                name = request.form.get('name')
                phone = request.form.get('phone')
                age = request.form.get('age')
                gender = request.form.get('gender')
                address = request.form.get('address')

                # Update user details except email
                update_data = {
                    "name": name,
                    "phone": phone,
                    "age": age,
                    "gender": gender,
                    "address": address
                }
                
                # Handle profile picture upload
                if 'profile_picture' in request.files:
                    profile_pic = request.files['profile_picture']
                    if profile_pic and profile_pic.filename != '':
                        # Check if it's a valid image
                        if profile_pic and allowed_file(profile_pic.filename):
                            # Read the image data
                            image_data = profile_pic.read()
                            
                            # Get file extension
                            file_ext = os.path.splitext(profile_pic.filename)[1].lower()
                            
                            # Update profile picture in database
                            update_data["profile_picture"] = {
                                "data": Binary(image_data),
                                "content_type": get_content_type(file_ext)
                            }
                        else:
                            flash("Invalid file format. Allowed formats: jpg, jpeg, png, gif", "error")
                            return redirect(url_for('profile'))

                mongo.db.users.update_one(
                    {"_id": ObjectId(user_id)},
                    {"$set": update_data}
                )
                flash("Profile updated successfully", "success")
                return redirect(url_for('profile'))
            
            elif 'change_password' in request.form:
                # Handle password change
                old_password = request.form.get('old_password')
                new_password = request.form.get('new_password')
                confirm_password = request.form.get('confirm_password')
                
                # Verify old password
                if not check_password_hash(user['password'], old_password):
                    flash("Current password is incorrect", "error")
                    return redirect(url_for('profile'))
                
                # Check if new passwords match
                if new_password != confirm_password:
                    flash("New passwords do not match", "error")
                    return redirect(url_for('profile'))
                
                # Update password
                hashed_password = generate_password_hash(new_password)
                mongo.db.users.update_one(
                    {"_id": ObjectId(user_id)},
                    {"$set": {"password": hashed_password}}
                )
                
                # Send email notification
                try:
                    email = user['email']
                    name = user['name']
                    
                    msg = Message(
                        "Password Change Notification - RapiACT",
                        sender=app.config['MAIL_USERNAME'],
                        recipients=[email]
                    )
                    
                    msg.html = f"""
                    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto; padding: 20px; border: 1px solid #ddd; border-radius: 5px;">
                        <div style="text-align: center; margin-bottom: 20px;">
                            <h2 style="color: #0056b3;">RapiACT Security Alert</h2>
                        </div>
                        <p>Hello {name},</p>
                        <p>This is to inform you that your RapiACT account password was changed successfully. This change was made on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}.</p>
                        <p>If you did not make this change, please contact our support team immediately or reset your password.</p>
                        <div style="margin-top: 30px; padding-top: 20px; border-top: 1px solid #ddd; font-size: 0.9em; color: #777;">
                            <p>This is an automated message. Please do not reply to this email.</p>
                            <p>RapiACT Team</p>
                        </div>
                    </div>
                    """
                    
                    mail.send(msg)
                except Exception as e:
                    # Log the error but don't prevent password change
                    print(f"Error sending password change email: {str(e)}")
                
                flash("Password changed successfully", "success")
                return redirect(url_for('profile'))

        return render_template('profile.html', user=user)
    
    except Exception as e:
        return f"Error fetching/updating user details: {str(e)}", 500

@app.route('/give_feedback/<appointment_id>', methods=['POST'])
def give_feedback(appointment_id):
    feedback = request.form.get('feedback')
    if not feedback:
        return jsonify({"error": "Feedback cannot be empty"}), 400

    try:
        obj_id = ObjectId(appointment_id)
    except Exception:
        return jsonify({"error": "Invalid appointment ID"}), 400

    result = db.appointments.update_one({"_id": obj_id}, {"$set": {"feedback": feedback}})
    if result.modified_count == 0:
        return jsonify({"error": "Feedback not saved. Appointment not found."}), 400

    flash("Feedback submitted successfully!", "success")
    return redirect(url_for('doctor_dashboard'))

@app.route('/mark_as_done/<appointment_id>', methods=['POST'])
def mark_as_done(appointment_id):
    try:
        obj_id = ObjectId(appointment_id)
    except Exception:
        return jsonify({"error": "Invalid appointment ID"}), 400

    result = db.appointments.update_one({"_id": obj_id}, {"$set": {"status": "done"}})
    if result.modified_count == 0:
        return jsonify({"error": "Could not mark appointment as done"}), 400

    flash("Appointment marked as done!", "success")
    return redirect(url_for('doctor_dashboard'))

@app.route('/get_report/<file_id>')
def get_report(file_id):
    try:
        file = fs.get(ObjectId(file_id))
        return Response(file.read(), mimetype=file.content_type, headers={"Content-Disposition": f"inline; filename={file.filename}"})
    except:
        flash("File not found.", "danger")
        return redirect(url_for('doctor_dashboard'))

# Creation of Appointment
@app.route('/create_appointment', methods=['GET', 'POST'])
def create_appointment():
    if 'user_id' in session and session.get('role') == 'patient':
        if request.method == 'POST':
            patient_name = request.form['patient_name']
            doctor_id = request.form['doctor_id']
            age = request.form['age']
            phone = request.form['phone']
            height = request.form['height']
            weight = request.form['weight']
            cause = request.form['cause']
            appointment_datetime = request.form['appointment_datetime']
            email = mongo.db.users.find_one({'_id': ObjectId(session['user_id'])})['email']

            # File upload (GridFS)
            report = request.files['report']
            file_id = None
            if report:
                file_id = fs.put(report, filename=secure_filename(report.filename), content_type=report.content_type)

            doctor = mongo.db.users.find_one({'_id': ObjectId(doctor_id)})
            doctor_name = f"{doctor['name']}"
            doctor_email = doctor['email']

            # Format datetime for display
            formatted_datetime = datetime.strptime(appointment_datetime, "%Y-%m-%dT%H:%M").strftime("%A, %B %d, %Y at %I:%M %p")

            # Insert appointment into DB
            appointment_id = mongo.db.appointments.insert_one({
                'patient_name': patient_name,
                'doctor_id': ObjectId(doctor_id),
                'age': age,
                'phone': phone,
                'email': email,
                'height': height,
                'weight': weight,
                'cause': cause,
                'appointment_datetime': appointment_datetime,
                'report_id': file_id,
                'patient_id': ObjectId(session['user_id']),
            }).inserted_id

            # Generate dashboard URLs for email templates
            dashboard_url = request.host_url.rstrip('/') + url_for('patient_dashboard')
            doctor_dashboard_url = request.host_url.rstrip('/') + url_for('doctor_dashboard')

            # HTML Email Template for Patient
            patient_email_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>Appointment Confirmation</title>
                <style>
                    body {{
                        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                        line-height: 1.6;
                        color: #333;
                        margin: 0;
                        padding: 0;
                    }}
                    .email-container {{
                        max-width: 600px;
                        margin: 0 auto;
                        padding: 20px;
                    }}
                    .header {{
                        background-color: #4285f4;
                        color: white;
                        padding: 20px;
                        text-align: center;
                        border-radius: 5px 5px 0 0;
                    }}
                    .content {{
                        background-color: #ffffff;
                        padding: 20px;
                        border-left: 1px solid #e0e0e0;
                        border-right: 1px solid #e0e0e0;
                    }}
                    .appointment-details {{
                        background-color: #f9f9f9;
                        padding: 15px;
                        border-radius: 5px;
                        margin: 15px 0;
                    }}
                    .appointment-item {{
                        margin-bottom: 10px;
                    }}
                    .appointment-label {{
                        font-weight: bold;
                        color: #4285f4;
                    }}
                    .footer {{
                        background-color: #f5f5f5;
                        padding: 15px;
                        text-align: center;
                        font-size: 12px;
                        color: #757575;
                        border-radius: 0 0 5px 5px;
                        border: 1px solid #e0e0e0;
                    }}
                    .button {{
                        display: inline-block;
                        background-color: #4285f4;
                        color: white;
                        padding: 10px 20px;
                        text-decoration: none;
                        border-radius: 5px;
                        margin-top: 15px;
                    }}
                    .logo {{
                        font-size: 24px;
                        font-weight: bold;
                        color: white;
                    }}
                </style>
            </head>
            <body>
                <div class="email-container">
                    <div class="header">
                        <div class="logo">RapiACT!</div>
                        <h2>Appointment Confirmation</h2>
                    </div>
                    <div class="content">
                        <p>Dear {patient_name},</p>
                        <p>Your appointment with Dr. {doctor_name} has been successfully scheduled.</p>
                        
                        <div class="appointment-details">
                            <div class="appointment-item">
                                <span class="appointment-label">Date & Time:</span> {formatted_datetime}
                            </div>
                            <div class="appointment-item">
                                <span class="appointment-label">Doctor:</span> Dr. {doctor_name}
                            </div>
                            <div class="appointment-item">
                                <span class="appointment-label">Reason:</span> {cause}
                            </div>
                        </div>
                        
                        <p>Please arrive 15 minutes before your scheduled appointment time.</p>
                        <p>You can view or manage your appointments from your patient dashboard.</p>
                        
                        <center>
                            <a href="{dashboard_url}" class="button">Access Dashboard</a>
                        </center>
                    </div>
                    <div class="footer">
                        <p>© 2025 RapiACT! Healthcare Services</p>
                        <p>This is an automated email. Please do not reply to this message.</p>
                    </div>
                </div>
            </body>
            </html>
            """

            # HTML Email Template for Doctor
            doctor_email_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>New Appointment Notification</title>
                <style>
                    body {{
                        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                        line-height: 1.6;
                        color: #333;
                        margin: 0;
                        padding: 0;
                    }}
                    .email-container {{
                        max-width: 600px;
                        margin: 0 auto;
                        padding: 20px;
                    }}
                    .header {{
                        background-color: #34a853;
                        color: white;
                        padding: 20px;
                        text-align: center;
                        border-radius: 5px 5px 0 0;
                    }}
                    .content {{
                        background-color: #ffffff;
                        padding: 20px;
                        border-left: 1px solid #e0e0e0;
                        border-right: 1px solid #e0e0e0;
                    }}
                    .patient-details {{
                        background-color: #f9f9f9;
                        padding: 15px;
                        border-radius: 5px;
                        margin: 15px 0;
                    }}
                    .patient-item {{
                        margin-bottom: 10px;
                    }}
                    .patient-label {{
                        font-weight: bold;
                        color: #34a853;
                    }}
                    .footer {{
                        background-color: #f5f5f5;
                        padding: 15px;
                        text-align: center;
                        font-size: 12px;
                        color: #757575;
                        border-radius: 0 0 5px 5px;
                        border: 1px solid #e0e0e0;
                    }}
                    .button {{
                        display: inline-block;
                        background-color: #34a853;
                        color: white;
                        padding: 10px 20px;
                        text-decoration: none;
                        border-radius: 5px;
                        margin-top: 15px;
                    }}
                    .logo {{
                        font-size: 24px;
                        font-weight: bold;
                        color: white;
                    }}
                </style>
            </head>
            <body>
                <div class="email-container">
                    <div class="header">
                        <div class="logo">RapiACT!</div>
                        <h2>New Appointment Notification</h2>
                    </div>
                    <div class="content">
                        <p>Dear Dr. {doctor_name},</p>
                        <p>You have a new appointment scheduled with a patient.</p>
                        
                        <div class="patient-details">
                            <div class="patient-item">
                                <span class="patient-label">Patient:</span> {patient_name}
                            </div>
                            <div class="patient-item">
                                <span class="patient-label">Date & Time:</span> {formatted_datetime}
                            </div>
                            <div class="patient-item">
                                <span class="patient-label">Age:</span> {age}
                            </div>
                            <div class="patient-item">
                                <span class="patient-label">Height:</span> {height}
                            </div>
                            <div class="patient-item">
                                <span class="patient-label">Weight:</span> {weight}
                            </div>
                            <div class="patient-item">
                                <span class="patient-label">Reason:</span> {cause}
                            </div>
                            <div class="patient-item">
                                <span class="patient-label">Contact:</span> {phone}
                            </div>
                            <div class="patient-item">
                                <span class="patient-label">Email:</span> {email}
                            </div>
                        </div>
                        
                        <p>Please review the appointment details and be prepared for the consultation.</p>
                        
                        <center>
                            <a href="{doctor_dashboard_url}" class="button">View Your Schedule</a>
                        </center>
                    </div>
                    <div class="footer">
                        <p>© 2025 RapiACT! Healthcare Services</p>
                        <p>This is an automated email. Please do not reply to this message.</p>
                    </div>
                </div>
            </body>
            </html>
            """

            # HTML Email Template for Reminder
            reminder_email_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>Appointment Reminder</title>
                <style>
                    body {{
                        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                        line-height: 1.6;
                        color: #333;
                        margin: 0;
                        padding: 0;
                    }}
                    .email-container {{
                        max-width: 600px;
                        margin: 0 auto;
                        padding: 20px;
                    }}
                    .header {{
                        background-color: #fbbc05;
                        color: white;
                        padding: 20px;
                        text-align: center;
                        border-radius: 5px 5px 0 0;
                    }}
                    .content {{
                        background-color: #ffffff;
                        padding: 20px;
                        border-left: 1px solid #e0e0e0;
                        border-right: 1px solid #e0e0e0;
                    }}
                    .reminder-details {{
                        background-color: #fffde7;
                        padding: 15px;
                        border-radius: 5px;
                        margin: 15px 0;
                        border-left: 4px solid #fbbc05;
                    }}
                    .reminder-item {{
                        margin-bottom: 10px;
                    }}
                    .reminder-label {{
                        font-weight: bold;
                        color: #fbbc05;
                    }}
                    .footer {{
                        background-color: #f5f5f5;
                        padding: 15px;
                        text-align: center;
                        font-size: 12px;
                        color: #757575;
                        border-radius: 0 0 5px 5px;
                        border: 1px solid #e0e0e0;
                    }}
                    .button {{
                        display: inline-block;
                        background-color: #fbbc05;
                        color: white;
                        padding: 10px 20px;
                        text-decoration: none;
                        border-radius: 5px;
                        margin-top: 15px;
                    }}
                    .logo {{
                        font-size: 24px;
                        font-weight: bold;
                        color: white;
                    }}
                    .countdown {{
                        font-size: 18px;
                        font-weight: bold;
                        color: #fbbc05;
                    }}
                </style>
            </head>
            <body>
                <div class="email-container">
                    <div class="header">
                        <div class="logo">RapiACT!</div>
                        <h2>Appointment Reminder</h2>
                    </div>
                    <div class="content">
                        <p>Dear {patient_name},</p>
                        <p><span class="countdown">Your appointment is in 5 minutes!</span></p>
                        
                        <div class="reminder-details">
                            <div class="reminder-item">
                                <span class="reminder-label">Date & Time:</span> {formatted_datetime}
                            </div>
                            <div class="reminder-item">
                                <span class="reminder-label">Doctor:</span> Dr. {doctor_name}
                            </div>
                            <div class="reminder-item">
                                <span class="reminder-label">Reason:</span> {cause}
                            </div>
                        </div>
                        
                        <p>Please ensure you're ready for your appointment.</p>
                        
                        <center>
                            <a href="{dashboard_url}" class="button">Access Dashboard</a>
                        </center>
                    </div>
                    <div class="footer">
                        <p>© 2025 RapiACT! Healthcare Services</p>
                        <p>This is an automated email. Please do not reply to this message.</p>
                    </div>
                </div>
            </body>
            </html>
            """

            # Send HTML emails
            subject = "Appointment Confirmation"
            send_email1(subject, [email], None, patient_email_html)
            send_email1(subject, [doctor_email], None, doctor_email_html)

            # Schedule reminder email
            appointment_time = datetime.strptime(appointment_datetime, "%Y-%m-%dT%H:%M")
            reminder_time = appointment_time - timedelta(minutes=5)
            reminder_subject = "Appointment Reminder"
            schedule_email1(reminder_subject, [email], None, reminder_email_html, reminder_time)

            flash(f"Appointment created successfully with Dr. {doctor_name}!", "success")
            return redirect(url_for('patient_dashboard'))

        doctors = list(mongo.db.users.find({'role': 'doctor'}))
        return render_template('create_appointment.html', doctors=doctors)

    flash("Unauthorized access.", "danger")
    return redirect(url_for('login'))

# Update send_email function to support HTML emails
def send_email1(subject, recipients, text_body=None, html_body=None):
    msg = Message(subject=subject, recipients=recipients)
    if text_body:
        msg.body = text_body
    if html_body:
        msg.html = html_body
    mail.send(msg)

# Update schedule_email function to support HTML emails
def schedule_email1(subject, recipients, text_body=None, html_body=None, scheduled_time=None):
    # This is a placeholder for scheduling emails
    # In a production app, you'd use a task queue like Celery with Redis/RabbitMQ
    # For simplicity, we'll simulate scheduling with a background thread
    
    def send_scheduled_email():
        # Calculate sleep time
        now = datetime.now()
        if scheduled_time > now:
            sleep_seconds = (scheduled_time - now).total_seconds()
            time.sleep(sleep_seconds)
        
        # Send email after sleeping
        send_email(subject, recipients, text_body, html_body)
    
    if scheduled_time:
        # Run in background thread
        email_thread = threading.Thread(target=send_scheduled_email)
        email_thread.daemon = True
        email_thread.start()
    else:
        # Send immediately if no scheduled time
        send_email(subject, recipients, text_body, html_body)

@app.route('/add_user', methods=['POST'])
def add_user():
    try:
        name = request.form.get('name')
        email = request.form.get('email')
        password = request.form.get('password')
        confirm_password = request.form.get('confirm_password')
        role = request.form.get('role')
        
        # Check if email already exists
        existing_user = mongo.db.users.find_one({"email": email})
        if existing_user:
            return jsonify({"success": False, "message": "Email already registered. Please use a different email."})
            
        if password != confirm_password:
            return jsonify({"success": False, "message": "Passwords do not match."})
        
        # Handle file upload
        photo_id = None
        photo = request.files.get('photo')
        if photo and allowed_file(photo.filename):
            photo_id = fs.put(photo, filename=photo.filename)
        elif photo and photo.filename != '':
            return jsonify({"success": False, "message": "Invalid file type. Only PDF, JPG, JPEG, or PNG files are allowed."})
        
        hashed_password = generate_password_hash(password)
        
        # Common user data
        user_data = {
            "_id": ObjectId(),
            "email": email,
            "password": hashed_password,
            "name": name,
            "role": role,
            "phone": request.form.get('phone'),
            "age": request.form.get('age'),
            "gender": request.form.get('gender'),
            "address": request.form.get('address'),
            "photo": photo_id,
            "created_at": datetime.utcnow()
        }
        
        if role == 'doctor':
            user_data.update({
                "specialization": request.form.get('specialization'),
                "years_experience": request.form.get('years_experience'),
                "professional_experience": request.form.get('professional_experience'),
                "account_status": "pending"
            })
        
        # Insert into MongoDB
        mongo.db.users.insert_one(user_data)
        
        # Send welcome email
        send_welcome_email(name, email, password, role)
        
        return jsonify({"success": True, "message": f"User {name} registered successfully!"})
    except Exception as e:
        return jsonify({"success": False, "message": str(e)})

# Email sending function
def send_welcome_email(name, email, password, role):
    try:
        # Configure email settings
        msg = Message(
            subject="Welcome to RapiACT!❤️ Healthcare System - Account Created",
            sender=app.config['MAIL_DEFAULT_SENDER'],
            recipients=[email]
        )
        
        # Customize message based on role
        role_specific_message = "book appointments and manage your healthcare needs" if role == "patient" else "manage patients and handle appointments"
        approval_message = "" if role == "patient" else "<p>Your doctor account is pending approval. An administrator will review your information and approve your account soon.</p>"
        
        # HTML email with styling - only send this version
        msg.html = f"""
<!DOCTYPE html>
<html>
<head>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            color: #333333;
            max-width: 600px;
            margin: 0 auto;
        }}
        .container {{
            padding: 20px;
            background-color: #f9f9f9;
            border-radius: 8px;
        }}
        .header {{
            background-color: #4a90e2;
            color: white;
            padding: 15px;
            text-align: center;
            border-radius: 8px 8px 0 0;
        }}
        .content {{
            background-color: white;
            padding: 20px;
            border-radius: 0 0 8px 8px;
            border: 1px solid #e0e0e0;
        }}
        .credentials {{
            background-color: #f5f5f5;
            padding: 15px;
            margin: 15px 0;
            border-left: 4px solid #4a90e2;
        }}
        .footer {{
            text-align: center;
            margin-top: 20px;
            color: #777;
            font-size: 14px;
        }}
        .highlight {{
            font-weight: bold;
            color: #4a90e2;
        }}
        .button {{
            display: inline-block;
            background-color: #4CAF50;
            color: white;
            padding: 12px 20px;
            text-decoration: none;
            border-radius: 4px;
            margin: 15px 0;
            font-weight: bold;
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h2>RapiACT! Healthcare System</h2>
        </div>
        <div class="content">
            <h3>Welcome, {name}!</h3>
            <p>Your account has been successfully created. You can now log in to your account and <span class="highlight">{role_specific_message}</span>.</p>
            
            <div class="credentials">
                <p><strong>Login Credentials:</strong></p>
                <p>Email: <span class="highlight">{email}</span></p>
                <p>Password: <span class="highlight">{password}</span></p>
            </div>
            
            {approval_message}
            
            <p>For security reasons, we recommend changing your password after your first login.</p>
            
            <a href="https://capstone-2025-37jt.onrender.com/login" class="button">Log In Now</a>
            
            <p>If you did not create this account, please contact our support team immediately.</p>
        </div>
        <div class="footer">
            <p>Best regards,<br>RapiACT! Team❤️</p>
            <p>&copy; 2025 RapiACT! Healthcare. All rights reserved.</p>
        </div>
    </div>
</body>
</html>
"""
        # Removed the plain text version (msg.body)
        
        # Send email
        mail.send(msg)
        return True
    except Exception as e:
        print(f"Error sending email: {str(e)}")
        # Continue with registration even if email fails
        return False

@app.route('/add_user', methods=['GET'])
def add_user_page():
    return render_template('add_user.html')

@app.route('/delete_user/<user_id>', methods=['POST'])
def delete_user(user_id):
    try:
        user = mongo.db.users.find_one({"_id": ObjectId(user_id)})
        if not user:
            return jsonify({"success": False, "message": "User not found."})
        
        mongo.db.users.delete_one({"_id": ObjectId(user_id)})
        return jsonify({"success": True, "message": "User deleted successfully!"})
    except Exception as e:
        return jsonify({"success": False, "message": str(e)})


def get_user_by_id(user_id):
    return users_collection.find_one({"_id": user_id})

def update_user(user_id, updated_data):
    users_collection.update_one({"_id": user_id}, {"$set": updated_data})

@app.route('/edit_user/<user_id>', methods=['GET', 'POST'])
def edit_user(user_id):
    user = mongo.db.users.find_one({"_id": ObjectId(user_id)})
    
    if not user:
        flash("User not found.", "danger")
        return redirect(url_for('users'))
    
    if request.method == 'POST':
        data = request.form
        phone = data.get("phone")
        age = data.get("age")
        gender = data.get("gender")
        address = data.get("address")
        
        updated_data = {
            "name": data.get("name"),
            "phone": phone,
            "age": age,
            "gender": gender,
            "address": address,
            "updated_at": datetime.utcnow()
        }
        
        if user["role"] == "doctor":
            updated_data.update({
                "specialization": data.get("specialization"),
                "years_experience": data.get("years_experience"),
                "professional_experience": data.get("professional_experience", "")
            })
        
        if 'photo' in request.files:
            photo = request.files['photo']
            if photo.filename:
                photo_id = fs.put(photo, filename=photo.filename)
                updated_data["photo"] = photo_id
        
        mongo.db.users.update_one({"_id": ObjectId(user_id)}, {"$set": updated_data})
        flash("User updated successfully!", "success")
        return redirect(url_for('users'))
    
    return render_template('edit_user.html', user=user)

from bson.json_util import dumps

@app.route('/get_doctors', methods=['GET'])
def get_doctors():
    doctors = list(mongo.db.users.find({"role": "doctor"}))
    return dumps(doctors), 200, {'Content-Type': 'application/json'}


@app.route("/all_doctors")
def all_doctors():
    return render_template("all_doctors.html")

@app.route("/get_visitors")
def get_visitors():
    # Ensure the visitors collection exists
    visitor_data = mongo.db.visitors.find_one({})
    
    if not visitor_data:
        mongo.db.visitors.insert_one({"count": 1})  # Initialize if not present
        count = 1
    else:
        mongo.db.visitors.update_one({}, {"$inc": {"count": 1}})
        count = visitor_data["count"] + 1  # Predict the updated count
    
    return jsonify({"count": count})

# Logout route
@app.route('/logout')
def logout():
    session.clear() 
    flash("You have been logged out.", "success")
    return redirect(url_for('home'))

import json
@app.route('/email_dashboard')
def email_dashboard():    
    # Get all users for the individual email tab
    users = list(mongo.db.users.find({}, {
        'name': 1, 
        'email': 1, 
        'role': 1
    }))
    
    return render_template('email_dashboard.html', users=users)

import hashlib
import time

# Send Bulk Email Route
@app.route('/send_bulk_email', methods=['POST'])
def send_bulk_email():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    recipient_group = request.form.get('recipient_group')
    subject = request.form.get('subject')
    message = request.form.get('message')
    image_position = request.form.get('image_position', 'middle')
    
    # Handle image upload if present
    image_id = None
    if 'email_image' in request.files and request.files['email_image'].filename != '':
        file = request.files['email_image']
        if file and allowed_file(file.filename):
            # Store image in MongoDB
            filename = secure_filename(file.filename)
            file_data = file.read()
            timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
            
            image_id = mongo.db.email_images.insert_one({
                'filename': filename,
                'data': Binary(file_data),
                'content_type': file.content_type,
                'uploaded_by': session['user_id'],
                'timestamp': datetime.now()
            }).inserted_id
    
    # Query for recipient emails based on the selected group
    if recipient_group == 'all':
        recipients = list(mongo.db.users.find({}, {'email': 1, 'name': 1}))
    elif recipient_group == 'doctors':
        recipients = list(mongo.db.users.find({'role': 'doctor'}, {'email': 1, 'name': 1}))
    elif recipient_group == 'patients':
        recipients = list(mongo.db.users.find({'role': 'patient'}, {'email': 1, 'name': 1}))
    else:
        flash('Invalid recipient group', 'error')
        return redirect(url_for('email_dashboard'))
    
    # Send emails using Flask-Mail
    sent_count = send_emails_with_flask_mail(recipients, subject, message, image_id, image_position)
    
    # Log the email campaign
    mongo.db.email_logs.insert_one({
        'sent_by': session['user_id'],
        'recipient_group': recipient_group,
        'subject': subject,
        'message': message,
        'has_image': image_id is not None,
        'image_id': image_id,
        'sent_count': sent_count,
        'timestamp': datetime.now()
    })
    
    flash(f'Successfully sent emails to {sent_count} recipients', 'success')
    return redirect(url_for('email_dashboard'))

# Send Individual Email Route
@app.route('/send_individual_email', methods=['POST'])
def send_individual_email():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    user_ids = json.loads(request.form.get('user_ids', '[]'))
    subject = request.form.get('subject')
    message = request.form.get('message')
    image_position = request.form.get('image_position', 'middle')
    
    # Handle image upload if present
    image_id = None
    if 'email_image' in request.files and request.files['email_image'].filename != '':
        file = request.files['email_image']
        if file and allowed_file(file.filename):
            # Store image in MongoDB
            filename = secure_filename(file.filename)
            file_data = file.read()
            
            image_id = mongo.db.email_images.insert_one({
                'filename': filename,
                'data': Binary(file_data),
                'content_type': file.content_type,
                'uploaded_by': session['user_id'],
                'timestamp': datetime.now()
            }).inserted_id
    
    if not user_ids:
        flash('No recipients selected', 'error')
        return redirect(url_for('email_dashboard'))
    
    # Convert IDs to ObjectId
    object_ids = [ObjectId(user_id) for user_id in user_ids]
    
    # Query for recipient emails
    recipients = list(mongo.db.users.find(
        {'_id': {'$in': object_ids}}, 
        {'email': 1, 'name': 1}
    ))
    
    # Send emails using Flask-Mail
    sent_count = send_emails_with_flask_mail(recipients, subject, message, image_id, image_position)
    
    # Log the individual emails
    mongo.db.email_logs.insert_one({
        'sent_by': session['user_id'],
        'recipient_ids': user_ids,
        'subject': subject,
        'message': message,
        'has_image': image_id is not None,
        'image_id': image_id,
        'sent_count': sent_count,
        'timestamp': datetime.now()
    })
    
    flash(f'Successfully sent emails to {sent_count} recipients', 'success')
    return redirect(url_for('email_dashboard'))

# Send Custom Email Route
@app.route('/send_custom_email', methods=['POST'])
def send_custom_email():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    custom_email = request.form.get('custom_email')
    subject = request.form.get('subject')
    message = request.form.get('message')
    image_position = request.form.get('image_position', 'middle')
    
    # Handle image upload if present
    image_id = None
    if 'email_image' in request.files and request.files['email_image'].filename != '':
        file = request.files['email_image']
        if file and allowed_file(file.filename):
            # Store image in MongoDB
            filename = secure_filename(file.filename)
            file_data = file.read()
            
            image_id = mongo.db.email_images.insert_one({
                'filename': filename,
                'data': Binary(file_data),
                'content_type': file.content_type,
                'uploaded_by': session['user_id'],
                'timestamp': datetime.now()
            }).inserted_id
    
    if not custom_email:
        flash('No recipient email provided', 'error')
        return redirect(url_for('email_dashboard'))
    
    # Create a recipient object similar to what we'd get from the database
    recipients = [{'email': custom_email, 'name': 'Recipient'}]
    
    # Send email using Flask-Mail
    sent_count = send_emails_with_flask_mail(recipients, subject, message, image_id, image_position)
    
    # Log the custom email
    mongo.db.email_logs.insert_one({
        'sent_by': session['user_id'],
        'recipient_email': custom_email,
        'subject': subject,
        'message': message,
        'has_image': image_id is not None,
        'image_id': image_id,
        'sent_count': sent_count,
        'timestamp': datetime.now()
    })
    
    flash('Email sent successfully', 'success')
    return redirect(url_for('email_dashboard'))

# Add a route to serve images directly from MongoDB with authentication
@app.route('/email_image/<image_id>', methods=['GET'])
def get_email_image(image_id):
    # Check if user is authenticated
    if 'user_id' not in session:
        return "Unauthorized", 401
        
    try:
        # Convert string ID to ObjectId
        image_oid = ObjectId(image_id)
        
        # Retrieve image from MongoDB
        image = mongo.db.email_images.find_one({'_id': image_oid})
        
        if not image:
            return "Image not found", 404
            
        # Create response with image binary data
        response = make_response(image['data'])
        response.headers.set('Content-Type', image['content_type'])
        response.headers.set('Content-Disposition', f'inline; filename={image["filename"]}')
        
        return response
    except Exception as e:
        print(f"Error retrieving image: {str(e)}")
        return "Error retrieving image", 500

# Helper function to send emails using Flask-Mail with HTML styling and image
def send_emails_with_flask_mail(recipients, subject, message, image_id=None, image_position='middle'):
    sent_count = 0
    app_url = request.host_url.rstrip('/')  # Get base URL without trailing slash

    for recipient in recipients:
        try:
            # Get recipient name or use a generic greeting
            recipient_name = recipient.get('name', 'Valued Patient')

            # Convert newlines in message to <br> for HTML formatting
            escaped_message = message.replace('\n', '<br>')
            
            # Create a unique token for this email that includes the recipient id to prevent sharing
            # This makes sure each recipient only sees their own image
            recipient_id = str(recipient.get('_id', 'unknown'))
            timestamp = int(time.time())
            email_token = hashlib.sha256(f"{recipient_id}:{image_id}:{timestamp}:{app.config['SECRET_KEY']}".encode()).hexdigest()
            
            # Prepare image HTML if an image was uploaded
            image_html = ""
            if image_id:
                # Create a URL for the image that includes authentication token
                image_url = f"{app_url}/email_image/{image_id}?token={email_token}&recipient={recipient_id}"
                image_html = f'<div style="text-align: center; margin: 20px 0;"><img src="{image_url}" alt="Email Image" style="max-width: 100%; height: auto; border-radius: 8px;"></div>'

            # Create HTML email with styling and image placement
            html_content = f"""
            <!DOCTYPE html>
            <html lang="en">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>{subject}</title>
            </head>
            <body style="margin: 0; padding: 0; font-family: Arial, sans-serif; line-height: 1.6; color: #333333; background-color: #f5f5f5;">
                <div style="max-width: 600px; margin: 0 auto; padding: 20px; background-color: #ffffff; border-radius: 8px; box-shadow: 0 2px 5px rgba(0,0,0,0.1);">
                    <div style="text-align: center; padding: 20px 0; border-bottom: 1px solid #eeeeee;">
                        <h1 style="margin: 0; color: #3b82f6; font-size: 24px;">RapiACT! Medical Center</h1>
                    </div>
                    
                    {'<!-- Image at top position -->' if image_position == 'top' else ''}
                    {image_html if image_position == 'top' else ''}
                    
                    <div style="padding: 20px 0;">
                        <p style="margin-top: 0; font-size: 16px; font-weight: bold;">Dear {recipient_name},</p>
                        
                        {'<!-- Image at middle position -->' if image_position == 'middle' else ''}
                        {image_html if image_position == 'middle' else ''}
                        
                        <div style="font-size: 16px; line-height: 1.6; color: #333333;">
                            {escaped_message}
                        </div>
                        
                        {'<!-- Image at bottom position -->' if image_position == 'bottom' else ''}
                        {image_html if image_position == 'bottom' else ''}
                    </div>
                    
                    <div style="padding: 20px 0; border-top: 1px solid #eeeeee; font-size: 14px; color: #666666;">
                        <p style="margin: 0;">Thank you for choosing our services.</p>
                        <p style="margin: 5px 0 0;">If you have any questions, please don't hesitate to contact us.</p>
                    </div>
                    
                    <div style="background-color: #f8f9fa; text-align: center; padding: 15px; border-radius: 0 0 8px 8px; font-size: 12px; color: #666666;">
                        <p style="margin: 0;">© 2025 RapiACT! Medical Center. All rights reserved.</p>
                        <p style="margin: 5px 0 0;">123 Healthcare Avenue, Medical District, MD 12345</p>
                        <p style="margin: 5px 0 0;">
                            <a href="#" style="color: #3b82f6; text-decoration: none;">Privacy Policy</a> | 
                            <a href="#" style="color: #3b82f6; text-decoration: none;">Unsubscribe</a>
                        </p>
                    </div>
                </div>
            </body>
            </html>
            """

            # Create a Flask-Mail message with HTML only
            msg = Message(
                subject=subject,
                recipients=[recipient['email']],
                html=html_content,
                sender=app.config['MAIL_DEFAULT_SENDER']
            )

            # Force HTML content subtype (helps in some mail clients)
            msg.content_subtype = "html"

            # Send the email
            mail.send(msg)
            sent_count += 1

        except Exception as e:
            print(f"Error sending to {recipient['email']}: {str(e)}")
            continue

    return sent_count

# API route to get email sending status
@app.route('/email_status', methods=['GET'])
def email_status():
    if 'user_id' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
    
    # Get the latest 5 email logs
    logs = list(mongo.db.email_logs.find().sort('timestamp', -1).limit(5))
    
    # Convert ObjectId to string for JSON serialization
    for log in logs:
        log['_id'] = str(log['_id'])
        log['sent_by'] = str(log['sent_by'])
        if 'image_id' in log and log['image_id']:
            log['image_id'] = str(log['image_id'])
        log['timestamp'] = log['timestamp'].strftime('%Y-%m-%d %H:%M:%S')
    
    return jsonify({'logs': logs})

# Other existing routes will remain unchanged
@app.route('/pharmacy')
@app.route('/pharmacy/')
def pharmacy():
    if 'user_id' not in session:
        flash("Please login to access the pharmacy.", "danger")
        return redirect(url_for('login'))
    
    user = mongo.db.users.find_one({'_id': ObjectId(session['user_id'])})
    
    # Handle search functionality
    search_query = request.args.get('search', '')
    if search_query:
        # Search in name and description fields
        search_filter = {
            "$or": [
                {"name": {"$regex": search_query, "$options": "i"}},
                {"description": {"$regex": search_query, "$options": "i"}}
            ]
        }
        products = list(mongo.db.pharmacy_products.find(search_filter))
    else:
        products = list(mongo.db.pharmacy_products.find())

    # Convert binary image data to base64 string for rendering in Jinja2
    for product in products:
        if product.get("image"):
            product["image"] = base64.b64encode(product["image"]).decode("utf-8")  # Convert to base64 string
    
    patient_prescriptions = []
    if session.get('role') == 'patient':
        patient_id = ObjectId(session['user_id'])
        patient_prescriptions = list(mongo.db.prescriptions.find({"patient_id": patient_id, "status": "recommended"}))
        
        for prescription in patient_prescriptions:
            product = mongo.db.pharmacy_products.find_one({"_id": prescription.get("product_id")})
            if product:
                prescription.update({
                    "product_name": product.get("name"),
                    "product_description": product.get("description"),
                    "product_image": base64.b64encode(product["image"]).decode("utf-8") if product.get("image") else "",
                    "product_price": product.get("price", 0)
                })
                
                # Check if payment is verified
                prescription["payment_verified"] = prescription.get("payment_status") == "verified"
            else:
                prescription.update({
                    "product_name": "Unknown Product",
                    "product_description": "No description available",
                    "product_image": "",
                    "product_price": 0,
                    "payment_verified": False
                })
    
    return render_template('pharmacy.html', products=products, prescriptions=patient_prescriptions, 
                          is_patient=(session.get('role') == 'patient'), user=user, search_query=search_query)



@app.route('/admin_dashboard/add_medicine', methods=['GET', 'POST'])
def add_medicine():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash("Unauthorized access.", "danger")
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        if 'image' in request.files:
            image_file = request.files['image']
            image_data = image_file.read()
        else:
            image_data = None
        
        medicine_data = {
            "name": request.form['name'],
            "description": request.form['description'],
            "price": float(request.form['price']),
            "stock": int(request.form['stock']),
            "image": image_data
        }
        mongo.db.pharmacy_products.insert_one(medicine_data)
        flash("Medicine added successfully!", "success")
        return redirect(url_for('pharmacy'))
    
    return render_template('add_medicine.html')


@app.route('/admin_dashboard/update_medicine/<medicine_id>', methods=['GET', 'POST'])
def update_medicine(medicine_id):
    if 'user_id' not in session or session.get('role') != 'admin':
        flash("Unauthorized access.", "danger")
        return redirect(url_for('login'))
    
    medicine = mongo.db.pharmacy_products.find_one({"_id": ObjectId(medicine_id)})
    
    if medicine and medicine.get("image"):
        medicine["image"] = base64.b64encode(medicine["image"]).decode("utf-8")
    
    if request.method == 'POST':
        update_data = {
            "name": request.form['name'],
            "description": request.form['description'],
            "price": float(request.form['price']),
            "stock": int(request.form['stock'])
        }
        
        if 'image' in request.files and request.files['image'].filename:
            image_file = request.files['image']
            update_data["image"] = image_file.read()
        
        mongo.db.pharmacy_products.update_one({"_id": ObjectId(medicine_id)}, {"$set": update_data})
        flash("Medicine updated successfully!", "success")
        return redirect(url_for('pharmacy'))
    
    return render_template('update_medicine.html', medicine=medicine)


@app.route('/doctor_dashboard/prescribe_medicine/<patient_id>', methods=['GET', 'POST'])
def prescribe_medicine(patient_id):
    if 'user_id' not in session or session.get('role') != 'doctor':
        flash("Unauthorized access.", "danger")
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        product_ids = request.form.getlist('product_ids[]')

        if not product_ids:
            flash("Please select at least one medicine.", "warning")
            # Fetch medicines with image processing for display
            medicines = list(mongo.db.pharmacy_products.find())
            for medicine in medicines:
                if 'image' in medicine and isinstance(medicine['image'], bytes):
                    medicine['image_base64'] = base64.b64encode(medicine['image']).decode('utf-8')
            return render_template('prescribe_medicine.html', medicines=medicines, patient_id=patient_id)

        dosage_instructions = request.form.get('dosage_instructions', '')
        duration = request.form.get('duration', '7')
        duration_unit = request.form.get('duration_unit', 'days')
        special_instructions = request.form.get('special_instructions', '')
        send_notification = 'send_notification' in request.form

        # Fetch full product details and prepare medicine list
        medicines_data = []
        for pid in product_ids:
            try:
                product = mongo.db.pharmacy_products.find_one({'_id': ObjectId(pid)})
                
                if product:
                    # Convert image to Base64 if it exists
                    image_data = ""
                    if 'image' in product and isinstance(product['image'], bytes):
                        image_data = base64.b64encode(product['image']).decode('utf-8')

                    medicines_data.append({
                        'id': str(product['_id']),
                        'name': product.get('name', 'Unknown'),
                        'description': product.get('description', 'No description available'),
                        'price': float(product.get('price', 0)),
                        'stock': product.get('stock', 0),
                        'image': image_data
                    })
            except Exception as e:
                print(f"Error fetching product {pid}: {str(e)}")
                traceback.print_exc()

        # Create prescription document with full product details
        prescription = {
            "patient_id": ObjectId(patient_id),
            "doctor_id": ObjectId(session['user_id']),
            "medicines": medicines_data,
            "dosage_instructions": dosage_instructions,
            "duration": duration,
            "duration_unit": duration_unit,
            "special_instructions": special_instructions,
            "status": "recommended",
            "date": datetime.utcnow()
        }

        # Insert the prescription into MongoDB
        result = mongo.db.prescriptions.insert_one(prescription)

        flash("Medicines prescribed successfully.", "success")
        return redirect(url_for('doctor_dashboard'))

    # GET request: Fetch available medicines
    medicines = list(mongo.db.pharmacy_products.find())
    
    # Process images for display in the template
    for medicine in medicines:
        if 'image' in medicine and isinstance(medicine['image'], bytes):
            medicine['image_base64'] = base64.b64encode(medicine['image']).decode('utf-8')
    
    return render_template('prescribe_medicine.html', medicines=medicines, patient_id=patient_id)

@app.route('/buy_product/<product_id>', methods=['GET'])
def buy_product(product_id):
    if 'user_id' not in session:
        flash("Please login to purchase products.", "danger")
        return redirect(url_for('login'))
    
    product = mongo.db.pharmacy_products.find_one({"_id": ObjectId(product_id)})
    
    if not product:
        flash("Product not found.", "danger")
        return redirect(url_for('pharmacy'))

    # Store purchase as a prescription-like record with full product details
    purchase = {
        "patient_id": ObjectId(session['user_id']),
        "medicines": [
            {
                "id": str(product["_id"]),
                "name": product.get("name", "Unknown"),
                "description": product.get("description", "No description available"),
                "price": float(product.get("price", 0)),
                "stock": int(product.get("stock", 0)),
                "image": base64.b64encode(product["image"]).decode("utf-8") if product.get("image") else None
            }
        ],
        "status": "direct_purchase",
        "date": datetime.utcnow()
    }
    
    purchase_id = mongo.db.prescriptions.insert_one(purchase).inserted_id

    # Path to the stored QR code in the static folder
    qr_code_path = url_for('static', filename='images/payement.jpg')

    return render_template(
        'payment.html',
        prescription_id=purchase_id,
        qr_code=qr_code_path,
        medicines=purchase["medicines"],  # Send medicines list
        total_price=purchase["medicines"][0]["price"]
    )




@app.route('/patient_dashboard/proceed_payment/<prescription_id>', methods=['GET'])
def proceed_payment(prescription_id):
    if 'user_id' not in session:
        flash("Please login to proceed with payment.", "danger")
        return redirect(url_for('login'))
    
    prescription = mongo.db.prescriptions.find_one({"_id": ObjectId(prescription_id)})
    
    if not prescription or str(prescription.get("patient_id")) != session['user_id']:
        flash("Invalid request.", "danger")
        return redirect(url_for('pharmacy'))
    
    # Check if it's a doctor prescription (multiple medicines) or direct purchase
    if "medicines" in prescription:
        # Get product details for all medicines in the prescription
        medicines = []
        total_price = 0
        
        for medicine_id in prescription.get("medicines", []):
            product = mongo.db.pharmacy_products.find_one({"_id": medicine_id})
            if product:
                if product.get("image"):
                    product["image"] = base64.b64encode(product["image"]).decode("utf-8")
                medicines.append(product)
                total_price += float(product.get("price", 0))
        
        # Path to the stored QR code in the static folder
        qr_code_path = url_for('static', filename='images/payement.jpg')

        return render_template('payment.html', 
                              prescription_id=prescription_id, 
                              qr_code=qr_code_path, 
                              medicines=medicines, 
                              total_price=total_price,
                              prescription=prescription)
    else:
        # Handle direct purchase (single product)
        product = mongo.db.pharmacy_products.find_one({"_id": prescription.get("product_id")})
        if product and product.get("image"):
            product["image"] = base64.b64encode(product["image"]).decode("utf-8")
        
        # Path to the stored QR code in the static folder
        qr_code_path = url_for('static', filename='images/payement.jpg')

        return render_template('payment.html', 
                              prescription_id=prescription_id, 
                              qr_code=qr_code_path, 
                              product=product)


@app.route('/patient_dashboard/submit_payment', methods=['POST'])
def submit_payment():
    if 'user_id' not in session or session.get('role') != 'patient':
        flash("Unauthorized access.", "danger")
        return redirect(url_for('login'))
    
    prescription_id = request.form['prescription_id']
    utr_id = request.form['utr_id']
    
    # Update prescription with payment details
    mongo.db.prescriptions.update_one(
        {"_id": ObjectId(prescription_id)}, 
        {"$set": {"payment_status": "paid", "utr_id": utr_id}}
    )
    
    flash("Payment submitted successfully!", "success")
    return redirect(url_for('pharmacy'))



import base64

@app.route('/admin_dashboard/verify_payment')
def verify_payment():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash("Unauthorized access.", "danger")
        return redirect(url_for('login'))

    pipeline = [
        {"$match": {"payment_status": "paid"}},
        {"$lookup": {  
            "from": "users",
            "localField": "patient_id",
            "foreignField": "_id",
            "as": "patient"
        }},
        {"$lookup": {  
            "from": "pharmacy_products",
            "localField": "medicines",
            "foreignField": "_id",
            "as": "products"
        }},
        {"$unwind": {"path": "$patient", "preserveNullAndEmptyArrays": True}},  
    ]

    payments = list(mongo.db.prescriptions.aggregate(pipeline))

    # Ensure 'products' is always a list and convert images
    for payment in payments:
        if not isinstance(payment.get("products"), list):
            payment["products"] = []  # Ensure it's a list if empty

        for product in payment["products"]:
            if isinstance(product, dict) and "image" in product and isinstance(product["image"], bytes):
                product["image"] = base64.b64encode(product["image"]).decode('utf-8')

    return render_template('verify_payment.html', payments=payments)


@app.route('/admin_dashboard/verify_payment_action/<payment_id>', methods=['POST'])
def verify_payment_action(payment_id):
    if 'user_id' not in session or session.get('role') != 'admin':
        flash("Unauthorized access.", "danger")
        return redirect(url_for('login'))

    action = request.form.get('action')

    # Fetch payment details with correct $lookup
    pipeline = [
        {"$match": {"_id": ObjectId(payment_id)}},
        {"$lookup": {
            "from": "users",
            "localField": "patient_id",
            "foreignField": "_id",
            "as": "patient"
        }},
        {"$unwind": {"path": "$patient", "preserveNullAndEmptyArrays": True}},
        
        # Extract medicine "id" values from the medicines list
        {"$addFields": {
            "medicine_ids": {
                "$map": {
                    "input": "$medicines",
                    "as": "med",
                    "in": {"$toObjectId": "$$med.id"}  # Convert string ID to ObjectId
                }
            }
        }},
        
        # Perform lookup on pharmacy_products using extracted medicine_ids
        {"$lookup": {
            "from": "pharmacy_products",
            "localField": "medicine_ids",
            "foreignField": "_id",
            "as": "products"
        }}
    ]

    payment_details = list(mongo.db.prescriptions.aggregate(pipeline))

    if not payment_details:
        flash("Payment record not found.", "danger")
        return redirect(url_for('verify_payment'))

    payment = payment_details[0]
    patient = payment.get('patient', {})

    # Ensure medicines are correctly formatted
    medicines = payment.get('products', [])
    if not isinstance(medicines, list):
        medicines = []

    # Debugging: Print fetched medicine list
    print(f"DEBUG: Fetched Medicines: {medicines}")

    recipient_email = patient.get("email")
    if not recipient_email:
        flash("User email not found, unable to send notification.", "warning")
        return redirect(url_for('verify_payment'))

    # Common CSS styles for both emails
    email_css = """
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 650px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f9f9f9;
        }
        .email-container {
            background-color: white;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0, 0, 0, 0.1);
            padding: 30px;
        }
        .header {
            text-align: center;
            padding-bottom: 20px;
            border-bottom: 1px solid #eee;
            margin-bottom: 20px;
        }
        .logo {
            font-size: 24px;
            font-weight: bold;
            color: #2c3e50;
        }
        h2 {
            color: #2c3e50;
            margin-top: 0;
        }
        h3 {
            color: #3498db;
        }
        table {
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
            box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1);
        }
        th, td {
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }
        th {
            background-color: #3498db;
            color: white;
            font-weight: 600;
        }
        tr:nth-child(even) {
            background-color: #f2f2f2;
        }
        tr:hover {
            background-color: #e9f7fe;
        }
        .total-row {
            font-weight: bold;
            background-color: #eaf6ff;
        }
        .footer {
            margin-top: 30px;
            padding-top: 20px;
            border-top: 1px solid #eee;
            text-align: center;
            font-size: 0.9em;
            color: #7f8c8d;
        }
        .button {
            display: inline-block;
            background-color: #3498db;
            color: white;
            padding: 10px 20px;
            text-decoration: none;
            border-radius: 4px;
            margin-top: 15px;
        }
        .contact-info {
            margin-top: 20px;
            font-size: 0.9em;
        }
    """

    if action == "verify":
        mongo.db.prescriptions.update_one({"_id": ObjectId(payment_id)}, {"$set": {"payment_status": "verified"}})
        flash("Payment verified successfully.", "success")

        # Construct email with medicine details
        subject = "Order Initiated - RapiACT"
        medicine_rows = "".join(
            f"<tr><td>{med.get('name', 'Unknown')}</td><td>{med.get('description', 'No description')}</td><td>₹{med.get('price', 0)}</td></tr>"
            for med in medicines if med
        )

        # Debugging: Print generated email table
        print(f"DEBUG: Email Medicine Rows:\n{medicine_rows}")

        # Ensure the table isn't empty
        if not medicine_rows:
            medicine_rows = "<tr><td colspan='3'>No medicines available</td></tr>"

        total_amount = sum(med.get('price', 0) for med in medicines)
        order_id = str(payment.get('_id', 'Unknown'))[:8]  # Use first 8 chars of payment ID

        body = f"""
        <html>
        <head>
            <style>
                {email_css}
            </style>
        </head>
        <body>
            <div class="email-container">
                <div class="header">
                    <div class="logo">RapiACT</div>
                    <p>Healthcare at your fingertips</p>
                </div>
                
                <h2>Order Confirmation</h2>
                <p>Dear {patient.get('name', 'User')},</p>
                <p>Your payment has been successfully verified, and your order has been initiated.</p>
                
                <div style="background-color: #e8f4fd; padding: 15px; border-radius: 6px; margin: 20px 0;">
                    <p style="margin: 0;"><strong>Order ID:</strong> #{order_id}</p>
                    <p style="margin: 8px 0 0;"><strong>Date:</strong> {datetime.now().strftime('%B %d, %Y')}</p>
                </div>
                
                <h3>Order Details:</h3>
                <table>
                    <tr>
                        <th>Product</th>
                        <th>Description</th>
                        <th>Price</th>
                    </tr>
                    {medicine_rows}
                    <tr class="total-row">
                        <th colspan="2" style="text-align: right;">Total:</th>
                        <td>₹{total_amount}</td>
                    </tr>
                </table>
                
                <p>Your order will be processed shortly. You will receive a notification when your order is ready for delivery or pickup.</p>
                
                <p>Thank you for choosing RapiACT for your healthcare needs!</p>
                
                <div class="footer">
                    <p>Best Regards,<br>The RapiACT Team</p>
                    <div class="contact-info">
                        <p>If you have any questions, please contact our support team</p>
                        <p>© 2025 RapiACT. All rights reserved.</p>
                    </div>
                </div>
            </div>
        </body>
        </html>
        """

        send_html_email(subject, [recipient_email], body)

    elif action == "reject":
        mongo.db.prescriptions.update_one({"_id": ObjectId(payment_id)}, {"$set": {"payment_status": "rejected"}})
        flash("Payment rejected.", "danger")

        subject = "Payment Rejected - RapiACT"
        body = f"""
        <html>
        <head>
            <style>
                {email_css}
            </style>
        </head>
        <body>
            <div class="email-container">
                <div class="header">
                    <div class="logo">RapiACT</div>
                    <p>Healthcare at your fingertips</p>
                </div>
                
                <h2>Payment Rejected</h2>
                <p>Dear {patient.get('name', 'User')},</p>
                
                <div style="background-color: #ffebee; padding: 15px; border-radius: 6px; margin: 20px 0;">
                    <p style="margin: 0;">Unfortunately, your payment has been rejected.</p>
                </div>
                
                <p>This could be due to one of the following reasons:</p>
                <ul style="padding-left: 20px;">
                    <li>Payment validation issue</li>
                    <li>Incomplete payment information</li>
                    <li>Transaction error with your payment provider</li>
                </ul>
                
                <p>Please contact our support team for more details and assistance in resolving this issue.</p>
                
                <a href="#" class="button">Contact Support</a>
                
                <div class="footer">
                    <p>Thank you for using RapiACT!</p>
                    <p>Best Regards,<br>The RapiACT Team</p>
                    <div class="contact-info">
                        <p>© 2025 RapiACT. All rights reserved.</p>
                    </div>
                </div>
            </div>
        </body>
        </html>
        """

        send_html_email(subject, [recipient_email], body)

    return redirect(url_for('verify_payment'))



def send_html_email(subject, recipients, html_body):
    msg = Message(subject, recipients=recipients)
    msg.html = html_body
    mail.send(msg)


@app.route('/admin_dashboard/orders')
def admin_orders():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash("Unauthorized access.", "danger")
        return redirect(url_for('login'))
    
    # Fetch all verified orders with patient and product details
    pipeline = [
        {"$match": {"payment_status": "verified"}},

        # Join with patient details
        {"$lookup": {
            "from": "users",
            "localField": "patient_id",
            "foreignField": "_id",
            "as": "patient"
        }},
        {"$unwind": {"path": "$patient", "preserveNullAndEmptyArrays": True}},

        # Extract medicine IDs from the `medicines` array
        {"$addFields": {
            "medicine_ids": {
                "$map": {
                    "input": "$medicines",
                    "as": "med",
                    "in": {"$toObjectId": "$$med.id"}  # Convert medicine "id" string to ObjectId
                }
            }
        }},

        # Lookup pharmacy products using extracted medicine_ids
        {"$lookup": {
            "from": "pharmacy_products",
            "localField": "medicine_ids",
            "foreignField": "_id",
            "as": "products"
        }}
    ]

    orders = list(mongo.db.prescriptions.aggregate(pipeline))

    return render_template('admin_orders.html', orders=orders)

@app.route('/admin_dashboard/update_order_status/<order_id>', methods=['POST'])
def update_order_status(order_id):
    if 'user_id' not in session or session.get('role') != 'admin':
        flash("Unauthorized access.", "danger")
        return redirect(url_for('login'))

    new_status = request.form.get('status')

    # Fetch the order
    order = mongo.db.prescriptions.find_one({"_id": ObjectId(order_id)})
    if not order:
        flash("Order not found.", "danger")
        return redirect(url_for('admin_orders'))

    patient = mongo.db.users.find_one({"_id": order.get("patient_id")})
    if not patient:
        flash("Patient details not found.", "danger")
        return redirect(url_for('admin_orders'))

    # Update order status
    mongo.db.prescriptions.update_one({"_id": ObjectId(order_id)}, {"$set": {"status": new_status}})

    # Send email notification based on new status
    subject, email_body = generate_status_email(patient, order, new_status)
    send_html_email(subject, [patient.get("email")], email_body)

    flash(f"Order status updated to '{new_status}' and email sent to patient.", "success")
    return redirect(url_for('admin_orders'))

def generate_status_email(patient, order, status):
    # Extract medicine details correctly from 'medicines' array
    order_items = "".join(
        f"<tr><td>{med.get('name', 'Unknown')}</td><td>{med.get('description', 'No description')}</td><td>₹{med.get('price', 0)}</td></tr>"
        for med in order.get("medicines", [])
    )

    # Calculate total price
    total_price = sum(med.get('price', 0) for med in order.get("medicines", []))
    
    # Generate a simple order ID from the MongoDB ObjectId
    order_id = str(order.get('_id', 'Unknown'))[:8]

    subject = f"Order Update - RapiACT [{status}]"

    # Enhanced CSS for a more professional email
    body = f"""
    <html>
    <head>
        <style>
            body {{
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                line-height: 1.6;
                color: #333;
                max-width: 650px;
                margin: 0 auto;
                padding: 20px;
                background-color: #f9f9f9;
            }}
            .email-container {{
                background-color: white;
                border-radius: 8px;
                box-shadow: 0 2px 10px rgba(0, 0, 0, 0.1);
                padding: 30px;
            }}
            .header {{
                text-align: center;
                padding-bottom: 20px;
                border-bottom: 1px solid #eee;
                margin-bottom: 20px;
            }}
            .logo {{
                font-size: 24px;
                font-weight: bold;
                color: #2c3e50;
            }}
            h2 {{
                color: #2c3e50;
                margin-top: 0;
            }}
            h3 {{
                color: #3498db;
            }}
            table {{
                border-collapse: collapse;
                width: 100%;
                margin: 20px 0;
                box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1);
            }}
            th, td {{
                border: 1px solid #ddd;
                padding: 12px;
                text-align: left;
            }}
            th {{
                background-color: #3498db;
                color: white;
                font-weight: 600;
            }}
            tr:nth-child(even) {{
                background-color: #f2f2f2;
            }}
            tr:hover {{
                background-color: #e9f7fe;
            }}
            .total-row {{
                font-weight: bold;
                background-color: #eaf6ff;
            }}
            .status-badge {{
                display: inline-block;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
                color: white;
                background-color: #3498db;
            }}
            .status-processing {{
                background-color: #3498db;
            }}
            .status-shipped {{
                background-color: #f39c12;
            }}
            .status-delivered {{
                background-color: #2ecc71;
            }}
            .status-canceled {{
                background-color: #e74c3c;
            }}
            .footer {{
                margin-top: 30px;
                padding-top: 20px;
                border-top: 1px solid #eee;
                text-align: center;
                font-size: 0.9em;
                color: #7f8c8d;
            }}
            .order-info {{
                background-color: #eaf6ff;
                padding: 15px;
                border-radius: 6px;
                margin: 20px 0;
            }}
        </style>
    </head>
    <body>
        <div class="email-container">
            <div class="header">
                <div class="logo">RapiACT</div>
                <p>Healthcare at your fingertips</p>
            </div>
            
            <h2>Order Update</h2>
            <p>Dear {patient.get('name', 'User')},</p>
            
            <p>Your order status has been updated to: 
                <span class="status-badge status-{status.lower()}">
                    {status.upper()}
                </span>
            </p>
            
            <div class="order-info">
                <p><strong>Order ID:</strong> #{order_id}</p>
                <p><strong>Date:</strong> {order.get('date', datetime.now()).strftime('%B %d, %Y')}</p>
            </div>
            
            <h3>Order Details:</h3>
            <table>
                <tr>
                    <th>Product</th>
                    <th>Description</th>
                    <th>Price</th>
                </tr>
                {order_items}
                <tr class="total-row">
                    <th colspan="2" style="text-align: right;">Total:</th>
                    <td>₹{total_price}</td>
                </tr>
            </table>
            
            <p>Thank you for choosing RapiACT for your healthcare needs!</p>
            
            <div class="footer">
                <p>Best Regards,<br>The RapiACT Team</p>
                <p>© 2025 RapiACT. All rights reserved.</p>
            </div>
        </div>
    </body>
    </html>
    """

    return subject, body


@app.route('/patient/orders')
def patient_orders():
    if 'user_id' not in session or session.get('role') != 'patient':
        flash("Unauthorized access.", "danger")
        return redirect(url_for('login'))
    
    patient_id = ObjectId(session['user_id'])

    pipeline = [
        {"$match": {"patient_id": patient_id, "payment_status": {"$exists": True}}},

        # Extract medicine IDs from the `medicines` array
        {"$addFields": {
            "medicine_ids": {
                "$map": {
                    "input": "$medicines",
                    "as": "med",
                    "in": {"$toObjectId": "$$med.id"}  # Convert medicine "id" string to ObjectId
                }
            }
        }},

        # Lookup pharmacy products using extracted medicine_ids
        {"$lookup": {
            "from": "pharmacy_products",
            "localField": "medicine_ids",
            "foreignField": "_id",
            "as": "products"
        }}
    ]

    orders = list(mongo.db.prescriptions.aggregate(pipeline))

    # Convert binary images to base64 for each product
    for order in orders:
        for product in order.get("products", []):  # Iterate over multiple products
            if "image" in product and product["image"]:
                product["image"] = base64.b64encode(product["image"]).decode('utf-8')

    return render_template('patient_orders.html', orders=orders)

@app.route('/cart/checkout', methods=['POST'])
def cart_checkout():
    if not session.get('logged_in') or session.get('role') != 'patient':
        return jsonify({'success': False, 'message': 'Please log in as a patient to checkout'}), 401
    
    # Get the cart items from the request
    data = request.json
    items = data.get('items', [])
    
    if not items:
        return jsonify({'success': False, 'message': 'Your cart is empty'}), 400
    
    try:
        # Calculate order details
        subtotal = sum(item['price'] * item['quantity'] for item in items)
        shipping_fee = 40.00  # Fixed shipping fee
        tax = subtotal * 0.05  # 5% GST
        total_amount = subtotal + shipping_fee + tax
        
        # Create order in database
        order_id = str(ObjectId())  # Generate a new ObjectId
        
        # Get the patient information
        patient_id = session.get('user_id')
        
        # Save order to database
        db.orders.insert_one({
            '_id': ObjectId(order_id),
            'patient_id': ObjectId(patient_id),
            'items': items,
            'subtotal': subtotal,
            'shipping_fee': shipping_fee,
            'tax': tax,
            'total_amount': total_amount,
            'status': 'pending_payment',
            'created_at': datetime.datetime.now(),
            'payment_details': {
                'status': 'pending'
            }
        })
        
        return jsonify({
            'success': True,
            'order_id': order_id,
            'message': 'Order created successfully'
        })
        
    except Exception as e:
        print(f"Error during checkout: {str(e)}")
        return jsonify({'success': False, 'message': 'An error occurred during checkout'}), 500


@app.route('/payment_cart/<order_id>')
def payment_cart(order_id):
    if not session.get('logged_in') or session.get('role') != 'patient':
        flash('Please log in as a patient to view this page', 'danger')
        return redirect(url_for('login'))
    
    try:
        # Get the order details from database
        order = db.orders.find_one({'_id': ObjectId(order_id)})
        
        if not order:
            flash('Order not found', 'danger')
            return redirect(url_for('pharmacy'))
        
        # Get patient details
        patient_id = session.get('user_id')
        patient = db.users.find_one({'_id': ObjectId(patient_id)})
        
        if not patient or str(order['patient_id']) != patient_id:
            flash('Unauthorized access', 'danger')
            return redirect(url_for('pharmacy'))
        
        # Generate a QR code for payment
        # This is a placeholder - you would typically generate this using a payment gateway
        qr_code_data = f"upi://pay?pa=merchant@upi&pn=RapiACT&am={order['total_amount']}&tr={order_id}&cu=INR"
        qr = qrcode.make(qr_code_data)
        
        # Create a BytesIO object to store the QR code
        qr_io = BytesIO()
        qr.save(qr_io)
        qr_io.seek(0)
        
        # Convert to base64 for displaying in HTML
        qr_base64 = base64.b64encode(qr_io.getvalue()).decode('utf-8')
        qr_code = f"data:image/png;base64,{qr_base64}"
        
        return render_template(
            'payment_cart.html',
            order=order,
            order_id=order_id,
            subtotal=order['subtotal'],
            shipping_fee=order['shipping_fee'],
            tax=order['tax'],
            total_amount=order['total_amount'],
            qr_code=qr_code
        )
        
    except Exception as e:
        print(f"Error during payment page loading: {str(e)}")
        flash('An error occurred while loading the payment page', 'danger')
        return redirect(url_for('pharmacy'))


@app.route('/submit_cart_payment', methods=['POST'])
def submit_cart_payment():
    if not session.get('logged_in') or session.get('role') != 'patient':
        return jsonify({'success': False, 'message': 'Please log in as a patient'}), 401
    
    data = request.json
    order_id = data.get('order_id')
    utr_id = data.get('utr_id')
    
    if not order_id or not utr_id:
        return jsonify({'success': False, 'message': 'Missing required fields'}), 400
    
    try:
        # Update the order with payment details
        result = db.orders.update_one(
            {'_id': ObjectId(order_id)},
            {'$set': {
                'payment_details': {
                    'utr_id': utr_id,
                    'status': 'verification_pending',
                    'submitted_at': datetime.datetime.now()
                },
                'status': 'payment_verification_pending'
            }}
        )
        
        if result.modified_count == 0:
            return jsonify({'success': False, 'message': 'Order not found or payment already submitted'}), 404
        
        return jsonify({
            'success': True,
            'message': 'Payment submitted successfully'
        })
        
    except Exception as e:
        print(f"Error during payment submission: {str(e)}")
        return jsonify({'success': False, 'message': 'An error occurred during payment submission'}), 500

import traceback
@app.route('/patient/prescriptions')
def patient_prescriptions():
    try:
        user_id = session.get('user_id')

        if not user_id:
            return jsonify({'error': 'User not logged in'}), 401

        # Fetch only prescriptions where status is "recommended"
        prescriptions = list(mongo.db.prescriptions.find({
            'patient_id': ObjectId(user_id),
            'status': 'recommended'  # Filter only recommended prescriptions
        }))

        formatted_prescriptions = []
        grand_total = 0  # To store the sum of all prescriptions' total cost

        for prescription in prescriptions:
            prescription['_id'] = str(prescription['_id'])

            # Retrieve medicine ObjectIds
            medicine_ids = prescription.get('medicines', [])

            # Fetch medicine details from pharmacy_products collection
            medicines = []
            total_cost = 0

            for medicine_id in medicine_ids:
                try:
                    if isinstance(medicine_id, ObjectId):  # If it's an ObjectId, fetch details
                        product = mongo.db.pharmacy_products.find_one({'_id': medicine_id})
                        if product:
                            medicine_data = {
                                'id': str(product['_id']),
                                'name': product.get('name', 'Unknown'),
                                'description': product.get('description', 'No description available'),
                                'price': float(product.get('price', 0)),
                                'stock': product.get('stock', 0),
                                'image': product.get('image', '')
                            }
                            medicines.append(medicine_data)
                            total_cost += medicine_data['price']
                    else:
                        medicines.append(medicine_id)  # If it's already a dict, use it directly
                        total_cost += medicine_id.get('price', 0)

                except Exception as e:
                    print(f"Error processing medicine {medicine_id}: {str(e)}")
                    traceback.print_exc()

            # Add to grand total
            grand_total += total_cost

            formatted_prescriptions.append({
                'prescription_id': prescription['_id'],
                'medicines': medicines,
                'total_cost': total_cost,
                'dosage_instructions': prescription.get('dosage_instructions', ''),
                'duration': prescription.get('duration', ''),
                'duration_unit': prescription.get('duration_unit', ''),
                'special_instructions': prescription.get('special_instructions', ''),
                'status': prescription.get('status', ''),
                'date': prescription.get('date', '')
            })

        # Render template or return JSON if requested
        if request.headers.get('Accept') == 'application/json':
            return jsonify({
                'prescriptions': formatted_prescriptions,
                'grand_total': grand_total
            })
        
        return render_template('prescriptions.html', prescriptions=formatted_prescriptions, grand_total=grand_total)

    except Exception as e:
        print(f"Error in patient_prescriptions API: {str(e)}")
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500



from flask import Flask, render_template, request, redirect, url_for
import joblib
import os
import re
import PyPDF2
from werkzeug.utils import secure_filename



from flask import Flask, render_template, request, redirect, url_for
import joblib
import PyPDF2
import re
import os
from werkzeug.utils import secure_filename


# Ensure uploads directory exists
if not os.path.exists('uploads'):
    os.makedirs('uploads')

# Load the model
model = joblib.load('templates/voting_diabetes .pkl')  # Note: fixed the space in the filename

# Features for diabetes prediction
features = [
    'Pregnancies', 'Glucose', 'Blood Pressure', 'Skin Thickness', 'Insulin',
    'BMI', 'Diabetes Pedigree Function',
]

def extract_text_from_pdf(filepath):
    text = ""
    with open(filepath, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + " "
            except UnicodeDecodeError:
                continue
    return text

def extract_data(text):
    extracted_data = {feature: None for feature in features}
    
    # Regex patterns with flexible spacing and colon handling
    patterns = {
        'Pregnancies': r'Pregnancies[:\s]*([\d\.]+)',
        'Glucose': r'Glucose[:\s]*([\d\.]+)',
        'Blood Pressure': r'Blood[\s]*Pressure[:\s]*([\d\.]+)',
        'Skin Thickness': r'Skin[\s]*Thickness[:\s]*([\d\.]+)',
        'Insulin': r'Insulin[:\s]*([\d\.]+)',
        'BMI': r'BMI[:\s]*([\d\.]+)',
        'Diabetes Pedigree Function': r'Diabetes[\s]*Pedigree[\s]*Function[:\s]*([\d\.]+)',
    }
    
    # Extract values using regex
    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            value = match.group(1)
            extracted_data[key] = float(value) if '.' in value else int(value)
        else:
            extracted_data[key] = 0  # Default to 0 if value isn't found
    
    print("✅ Extracted Data:", extracted_data)
    return extracted_data

# Define index route
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/diabetes', methods=['GET', 'POST'])
def diabetes():
    return render_template('diabetes.html', features=features, extracted_data=None)

@app.route('/extract', methods=['POST'])
def extract():
    file = request.files['file']
    if file:
        filename = secure_filename(file.filename)
        filepath = os.path.join('uploads', filename)
        file.save(filepath)
        
        if file.filename.endswith('.pdf'):
            text = extract_text_from_pdf(filepath)
        else:
            return "❌ Unsupported file format. Please upload a PDF."
        
        extracted_data = extract_data(text)
        return render_template('diabetes.html', features=features, extracted_data=extracted_data)
    return redirect(url_for('diabetes'))

@app.route('/predict', methods=['POST'])
def predict():
    input_data = []
    for feature in features:
        value = request.form.get(feature)
        input_data.append(float(value) if value else 0)
    
    # Predict using the loaded model
    prediction = model.predict([input_data])[0]
    prediction_text = '⚠️ High Risk of Diabetes' if prediction == 1 else '✅ Low Risk of Diabetes'
    
    # Suggest precautions if high risk
    precautions = [
        '🟢 Follow a balanced diet low in sugar.',
        '🏃 Exercise regularly.',
        '💉 Monitor blood sugar levels frequently.',
        '🩺 Regular check-ups with your doctor.'
    ] if prediction == 1 else []
    
    return render_template('diabetes.html', features=features, extracted_data=None, prediction_text=prediction_text, precautions=precautions)




































#Heart Disease
model2 = joblib.load('templates/heart.pkl')

# Descriptive features
features_heart = [
    'Age', 'Sex', 'Chest Pain Type', 'Resting Blood Pressure', 'Serum Cholesterol',
    'Fasting Blood Sugar', 'Resting ECG', 'Max Heart Rate Achieved', 'Exercise Induced Angina',
    'ST Depression', 'Slope of Peak Exercise', 'Number of Major Vessels Colored', 'Thalassemia'
]

# Categorical mappings
mappings_heart = {
    'Sex': {'Male': 0, 'Female': 1},
    'Fasting Blood Sugar': {'Normal': 0, 'High': 1},
    'Exercise Induced Angina': {'No': 0, 'Yes': 1}
}


def extract_data_heart(text):
    extracted_data = {feature: None for feature in features_heart}

    # Patterns to extract numerical and categorical data
    patterns_heart = {
        'Age': r'Age:\s*(\d+)',
        'Sex': r'Sex:\s*(Male|Female)',
        'Chest Pain Type': r'Chest Pain Type:\s*(\d+)',
        'Resting Blood Pressure': r'Resting Blood Pressure:\s*(\d+)',
        'Serum Cholesterol': r'Serum Cholesterol:\s*(\d+)',
        'Fasting Blood Sugar': r'Fasting Blood Sugar:\s*(Normal|High)',
        'Resting ECG': r'Resting ECG:\s*(\d+)',
        'Max Heart Rate Achieved': r'Max Heart Rate Achieved:\s*(\d+)',
        'Exercise Induced Angina': r'Exercise Induced Angina:\s*(Yes|No)',
        'ST Depression': r'ST Depression:\s*([\d\.]+)',
        'Slope of Peak Exercise': r'Slope of Peak Exercise:\s*(\d+)',
        'Number of Major Vessels Colored': r'Number of Major Vessels Colored:\s*(\d+)',
        'Thalassemia': r'Thalassemia:\s*(\d+)'  
    }

    # Extract values using regex patterns
    for key, pattern in patterns_heart.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            value = match.group(1)
            # Convert categorical values to numeric
            if key in mappings_heart:
                extracted_data[key] = mappings_heart[key].get(value, None)
            else:
                extracted_data[key] = float(value)
    
    print("✅ Extracted Data:", extracted_data)
    return extracted_data

@app.route('/model2', methods=['GET', 'POST'])
def heart_model1():
    return render_template('heart_home.html', features=features_heart, extracted_data=None)

@app.route('/extract1', methods=['POST'])
def extract_heart():
    file = request.files['file']
    if file:
        filename = secure_filename(file.filename)
        filepath = os.path.join('uploads', filename)
        file.save(filepath)

        if file.filename.endswith('.pdf'):
            text = extract_text_from_pdf(filepath)
        else:
            return "❌ Unsupported file format. Please upload a PDF."
        
        extracted_data = extract_data_heart(text)
        return render_template('heart_home.html', features=features_heart, extracted_data=extracted_data)
    return redirect(url_for('model2'))

@app.route('/predict/heart', methods=['POST'])
def predict_heart():
    input_data = []
    for feature in features_heart:
        value = request.form.get(feature)
        if feature in ['Sex', 'Fasting Blood Sugar', 'Exercise Induced Angina']:
            value = 1 if value and value.lower() in ['female', 'high', 'yes'] else 0
        input_data.append(float(value) if value else 0)
    
    # Predict using the loaded model
    prediction = model2.predict([input_data])[0]
    prediction_text = '⚠️ High Risk of Heart Disease' if prediction == 1 else '✅ Low Risk of Heart Disease'

    # Suggest precautions if high risk
    precautions = [
        '🟢 Maintain a healthy diet.',
        '🏃 Exercise regularly.',
        '🧂 Monitor cholesterol levels.',
        '🩺 Regular heart check-ups.'
    ] if prediction == 1 else []

    return render_template('heart_home.html', features=features_heart, extracted_data=None, prediction_text=prediction_text, precautions=precautions)





# Load the trained model (Ensure the model file is in the same directory)
model3 = joblib.load('templates/kidney.pkl')

# List of features expected from the form
features_kidney = [
    'age', 'blood_pressure', 'specific_gravity', 'albumin', 'sugar',
    'red_blood_cells', 'pus_cell', 'pus_cell_clumps', 'bacteria',
    'blood_glucose_random', 'blood_urea', 'serum_creatinine', 'sodium',
    'potassium', 'haemoglobin', 'packed_cell_volume', 'white_blood_cell_count',
    'red_blood_cell_count', 'hypertension', 'diabetes_mellitus',
    'coronary_artery_disease', 'appetite', 'peda_edema', 'aanemia'
]
def extract_text_from_pdf_kidney(file):
    reader = PyPDF2.PdfReader(file)
    text = ''
    for page in reader.pages:
        text += page.extract_text() + ' '
    return text
# Function to extract text from DOCX
def extract_text_from_docx(file):
    doc = docx.Document(file)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + ' '
    return text

# Function to extract data from text using regex
def extract_data_from_kidney(text):
    extracted_data = {}
    patterns_kidney = {
        'age': r'Age:\s*(\d+)',
        'blood_pressure': r'Blood Pressure:\s*(\d+)',
        'specific_gravity': r'Specific Gravity:\s*([\d.]+)',
        'albumin': r'Albumin:\s*(\d+)',
        'sugar': r'Sugar:\s*(\d+)',
        'red_blood_cells': r'Red Blood Cells:\s*(abnormal|normal)',
        'pus_cell': r'Pus Cell:\s*(abnormal|normal)',
        'pus_cell_clumps': r'Pus Cell Clumps:\s*(present|not present|absent)',
        'bacteria': r'Bacteria:\s*(present|not present|absent)',
        'blood_glucose_random': r'Blood Glucose Random:\s*(\d+)',
        'blood_urea': r'Blood Urea:\s*(\d+)',
        'serum_creatinine': r'Serum Creatinine:\s*([\d.]+)',
        'sodium': r'Sodium:\s*(\d+)',
        'potassium': r'Potassium:\s*([\d.]+)',
        'haemoglobin': r'Haemoglobin:\s*([\d.]+)',
        'packed_cell_volume': r'Packed Cell Volume:\s*(\d+)',
        'white_blood_cell_count': r'White Blood Cell Count:\s*(\d+)',
        'red_blood_cell_count': r'Red Blood Cell Count:\s*([\d.]+)',
        'hypertension': r'Hypertension:\s*(yes|no)',
        'diabetes_mellitus': r'Diabetes Mellitus:\s*(yes|no)',
        'coronary_artery_disease': r'Coronary Artery Disease:\s*(yes|no)',
        'appetite': r'Appetite:\s*(good|poor)',
        'peda_edema': r'Peda Edema:\s*(yes|no)',
        'aanemia': r'Aanemia:\s*(yes|no)'
    }

    for feature, pattern in patterns_kidney.items():
        match = re.search(pattern, text, re.IGNORECASE)
        extracted_data[feature] = match.group(1) if match else None

    return extracted_data


# Home Route
@app.route('/model3', methods=['GET'])
def kidney_model():
    return render_template('kidney_home.html', features=features_kidney, extracted_data={}, prediction_text=None, precautions=None)

# File Upload Route
@app.route('/extract/kidney', methods=['POST'])
def extract_kidney():
    file = request.files['file']
    if file.filename.endswith('.pdf'):
        text = extract_text_from_pdf_kidney(file)
    elif file.filename.endswith('.docx'):
        text = extract_text_from_docx(file)
    else:
        return "Unsupported file format. Please upload a PDF or DOCX."

    extracted_data = extract_data_from_kidney(text)
    return render_template('kidney_home.html', features=features_kidney, extracted_data=extracted_data, prediction_text=None, precautions=None)

# Prediction Route (Using Trained Model)
@app.route('/predict/kidney', methods=['POST'])
def predict_kidney():
    input_data = {feature: request.form.get(feature) for feature in features_kidney}

    # Convert input data for the model
    processed_data = []
    for feature in features_kidney:
        value = input_data.get(feature)
        if value.lower() in ['abnormal', 'present', 'yes', 'poor']:
            processed_data.append(1)
        elif value.lower() in ['normal', 'not present', 'no', 'good', 'absent']:
            processed_data.append(0)
        else:
            try:
                processed_data.append(float(value))
            except (TypeError, ValueError):
                processed_data.append(0.0)

    # Model prediction
    prediction = model3.predict([processed_data])[0]

    # Prediction result and precautions
    prediction_text = "Positive (Kidney Disease Detected)" if prediction == 0 else "Negative (No Kidney Disease)"
    precautions = [
        "Maintain a healthy diet low in sodium.",
        "Monitor blood pressure regularly.",
        "Stay hydrated and avoid smoking."
    ] if prediction == 0 else ["No immediate precautions needed. Continue regular health check-ups."]

    return render_template('kidney_home.html', features=features_kidney, extracted_data=input_data, prediction_text=prediction_text, precautions=precautions)


import pickle
import numpy as np
import os
import re
from flask import Flask, render_template, request, redirect, url_for, flash, session
import pdfplumber
import docx
from PyPDF2 import PdfReader
from docx import Document



# Configure separate upload folders for each module
PARKINSONS_UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'parkinsons_uploads')
THYROID_UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'thyroid_uploads')

# Create the upload directories if they don't exist
if not os.path.exists(PARKINSONS_UPLOAD_FOLDER):
    os.makedirs(PARKINSONS_UPLOAD_FOLDER)
if not os.path.exists(THYROID_UPLOAD_FOLDER):
    os.makedirs(THYROID_UPLOAD_FOLDER)

# Load pre-trained models and scalers
parkinsons_model = pickle.load(open('templates/svmodel.pkl', 'rb'))
parkinsons_scaler = pickle.load(open('templates/scaling.pkl', 'rb'))

thyroid_model = pickle.load(open('templates/best_model1.pkl', 'rb'))
thyroid_scaler = pickle.load(open('templates/scaler1.pkl', 'rb'))

# Keywords for matching
KEYWORDS_PARKINSONS = [
    "MDVP:Fo(Hz)", "MDVP:Fhi(Hz)", "MDVP:Flo(Hz)", "MDVP:Jitter(%)",
    "MDVP:Jitter(Abs)", "MDVP:RAP", "MDVP:PPQ", "Jitter:DDP", "MDVP:Shimmer",
    "MDVP:Shimmer(dB)", "Shimmer:APQ3", "Shimmer:APQ5", "MDVP:APQ", "Shimmer:DDA",
    "NHR", "HNR", "RPDE", "DFA", "spread1", "spread2", "D2", "PPE"
]

THYROID_KEYWORDS = [
    "Age", "Gender", "Smoking", "Smoking History",
    "Radiotherapy History", "Thyroid Function",
    "Physical Examination", "Adenopathy",
    "Types of Thyroid Cancer (Pathology)", "Focality",
    "Risk", "Tumor", "Lymph Nodes",
    "Cancer Metastasis", "Stage", "Treatment Response"
]

# Utility functions for text extraction
def extract_text_from_pdf(file):
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        try:
            page_text = page.extract_text()
            if page_text:
                cleaned_text = re.sub(r"[•\t\r\u200B]", " ", page_text).strip()
                text += cleaned_text + "\n"
        except UnicodeEncodeError:
            continue
    return text.strip()

def extract_text_from_docx(file):
    doc = Document(file)
    text = ""
    for para in doc.paragraphs:
        try:
            cleaned_text = re.sub(r"[•\t\r\u200B]", " ", para.text).strip()
            text += cleaned_text + "\n"
        except UnicodeEncodeError:
            continue
    return text.strip()

#--------- PARKINSONS DISEASE MODULE ---------#

# Parkinsons home page route
@app.route('/parkinsons')
def parkinsons():
    return render_template('parkinsons.html', features_parkinsons=KEYWORDS_PARKINSONS, extracted_data={})

# Parkinsons file upload and extraction route 
@app.route('/extract/parkinsons', methods=['POST'])
def extract_data_parkinsons():
    extracted_data = {}
    status_message = None

    # File Upload Handling
    if 'file' in request.files and request.files['file'].filename != '':
        file = request.files['file']
        safe_filename = file.filename.encode('utf-8', errors='ignore').decode()
        filename = os.path.join(PARKINSONS_UPLOAD_FOLDER, safe_filename)
        file.save(filename)

        # Detect file type and extract text
        if file.filename.endswith('.pdf'):
            text = extract_text_from_pdf(filename)
            status_message = f"PDF processed: {safe_filename}"
        elif file.filename.endswith('.docx'):
            text = extract_text_from_docx(filename)
            status_message = f"DOCX processed: {safe_filename}"
        else:
            return render_template('parkinsons.html', 
                                  error='Unsupported file format', 
                                  features_parkinsons=KEYWORDS_PARKINSONS,
                                  extracted_data={})

        # Extract data using regular expressions
        extracted_data = regex_extract_parkinsons(text)

    return render_template(
        "parkinsons.html",
        extracted_data=extracted_data,
        features_parkinsons=KEYWORDS_PARKINSONS,
        status_message=status_message
    )

# Parkinsons prediction route
@app.route('/predict/parkinsons', methods=['POST'])
def predict_parkinsons():
    input_data = []
    extracted_data = {}

    # Collect data from form fields (manual or auto-filled)
    for keyword in KEYWORDS_PARKINSONS:
        value = request.form.get(keyword)
        try:
            if value and value.strip() != '':
                input_data.append(float(value))
                extracted_data[keyword] = float(value)
            else:
                input_data.append(0.0)
        except (ValueError, TypeError):
            input_data.append(0.0)

    # Scale and predict
    input_scaled = parkinsons_scaler.transform(np.array(input_data).reshape(1, -1))
    prediction = parkinsons_model.predict(input_scaled)[0]

    # Prediction result
    result = "Positive for Parkinson's Disease" if prediction == 1 else "Negative for Parkinson's Disease"

    # Precautions for positive result
    precautions = []
    if prediction == 1:
        precautions = [
            "🧘 Engage in regular physical exercise and therapy.",
            "🥗 Follow a balanced diet rich in fiber and antioxidants.",
            "💊 Take medications as prescribed by your doctor.",
            "🩺 Schedule regular neurological checkups.",
            "🧠 Practice mental exercises and stress management."
        ]

    return render_template(
        "parkinsons.html",
        prediction_text=f"Prediction: {result}",
        precautions=precautions if prediction == 1 else None,
        extracted_data=extracted_data,
        features_parkinsons=KEYWORDS_PARKINSONS
    )

# Parkinsons regex-based data extraction
def regex_extract_parkinsons(text):
    extracted_data = {}

    # Clean hidden characters, normalize spaces, and handle special formatting
    text = re.sub(r"[•\t\r\u200B\u00A0]", " ", text)  # Remove unwanted characters
    text = re.sub(r"\s+", " ", text)  # Normalize spaces
    text = text.replace("–", "-")  # Replace special minus signs with standard minus

    # Regex pattern for accurate matching of negative, decimal, or scientific values
    pattern = re.compile(
        r"(MDVP:Fo\(Hz\)|MDVP:Fhi\(Hz\)|MDVP:Flo\(Hz\)|MDVP:Jitter\(%\)|"
        r"MDVP:Jitter\(Abs\)|MDVP:RAP|MDVP:PPQ|Jitter:DDP|MDVP:Shimmer|MDVP:Shimmer\(dB\)|"
        r"Shimmer:APQ3|Shimmer:APQ5|MDVP:APQ|Shimmer:DDA|NHR|HNR|RPDE|DFA|Spread1|Spread2|D2|PPE)"
        r"[:\s]*"  # Matches colon or spaces
        r"([-+]?\d*\.?\d+(?:[eE][-+]?\d+)?)",  # Captures negative, decimal, or scientific notation
        flags=re.IGNORECASE
    )

    matches = pattern.findall(text)

    # Log extracted values for debugging
    print("🔍 Parkinsons Extracted Matches:", matches)

    # Store extracted values in the dictionary
    for match in matches:
        keyword, value = match
        # Handle case-sensitivity in the keys
        for k in KEYWORDS_PARKINSONS:
            if k.lower() == keyword.lower():
                try:
                    extracted_data[k] = float(value)
                except ValueError:
                    extracted_data[k] = None
                break

    # Ensure all keywords are present
    for keyword in KEYWORDS_PARKINSONS:
        if keyword not in extracted_data:
            extracted_data[keyword] = None

    # Log extracted data for debugging
    print("✅ Parkinsons Final Extracted Data:", extracted_data)

    return extracted_data

#--------- THYROID MODULE ---------#

@app.route('/Tyroid')
def Tyroid():
    return render_template('tyroid.html', keywords=THYROID_KEYWORDS, extracted_data={})

@app.route('/extract/Tyroid', methods=['POST'])
def extract_data_thyroid():
    extracted_data = {}
    if 'file' in request.files and request.files['file'].filename != '':
        file = request.files['file']
        safe_filename = secure_filename(file.filename)
        filename = os.path.join(THYROID_UPLOAD_FOLDER, safe_filename)
        file.save(filename)

        try:
            if file.filename.endswith('.pdf'):
                text = extract_text_from_pdf(filename)
                status_message = f"PDF processed: {safe_filename}"
            elif file.filename.endswith('.docx'):
                text = extract_text_from_docx(filename)
                status_message = f"DOCX processed: {safe_filename}"
            else:
                return render_template('tyroid.html', error='Unsupported file format', 
                                      keywords=THYROID_KEYWORDS, extracted_data={})
            
            # Extract data from the document text
            extracted_data = regex_extract_thyroid(text)
            
            # Ensure all expected fields are present with proper formatting
            for keyword in THYROID_KEYWORDS:
                if keyword not in extracted_data:
                    extracted_data[keyword] = 0
                    
            print("✅ Thyroid Data prepared for form:", extracted_data)
            
        except Exception as e:
            print(f"Error during extraction: {str(e)}")
            return render_template('tyroid.html', error=f'Error extracting data: {str(e)}', 
                                  keywords=THYROID_KEYWORDS, extracted_data={})

    # Pass the extracted data to the template for autofilling
    return render_template("tyroid.html", extracted_data=extracted_data, 
                          keywords=THYROID_KEYWORDS, status_message=status_message)

@app.route('/predict/Tyroid', methods=['POST'])
def predict_thyroid():
    try:
        data = {}
        extracted_data = {}

        for keyword in THYROID_KEYWORDS:
            value = request.form.get(keyword)
            try:
                if value and value.strip() != '':
                    if keyword == "Age":
                        raw_age = float(value)
                        extracted_data[keyword] = raw_age
                        data[keyword] = thyroid_scaler.transform([[raw_age]])[0][0]
                    else:
                        data[keyword] = int(value)
                        extracted_data[keyword] = int(value)
                else:
                    data[keyword] = 0
                    extracted_data[keyword] = 0
            except (ValueError, TypeError):
                data[keyword] = 0
                extracted_data[keyword] = 0

        df_input = pd.DataFrame([data])
        prediction = thyroid_model.predict(df_input)[0]
        result = "High Risk of Thyroid Recurrence" if prediction == 1 else "Low Risk of Thyroid Recurrence"

        recommendations = [
            "🏥 Schedule more frequent follow-up appointments with your endocrinologist.",
            "🔬 Consider additional diagnostic imaging or blood tests.",
            "💊 Adhere strictly to your prescribed treatment.",
            "🍎 Maintain a healthy diet low in iodine if advised.",
            "📝 Keep a symptom diary."
        ] if prediction == 1 else [
            "🏥 Continue regular check-ups as advised.",
            "🩸 Monitor thyroid function tests regularly.",
            "🍎 Eat healthy and maintain lifestyle.",
            "📝 Report any new symptoms."
        ]

        return render_template("tyroid.html", prediction_text=f"Prediction: {result}", 
                              recommendations=recommendations, extracted_data=extracted_data, 
                              keywords=THYROID_KEYWORDS)
    except Exception as e:
        return render_template('tyroid.html', error=f"Prediction Error: {str(e)}", 
                              keywords=THYROID_KEYWORDS, extracted_data={})

def regex_extract_thyroid(text):
    extracted_data = {}
    
    # Normalize text for better pattern matching
    text = re.sub(r"[•\t\r\u200B\u00A0]", " ", text)
    text = re.sub(r"\s+", " ", text).lower()
    
    # Helper function to add keys with default values if none found
    def add_or_default(key, value, default=0):
        extracted_data[key] = value if value is not None else default

    # Age extraction with better pattern matching
    age_match = re.search(r"age[:\s\-]*([0-9]{1,3})(?:\s*years)?", text)
    add_or_default("Age", float(age_match.group(1)) if age_match else None)

    # Gender extraction (improved pattern matching)
    gender_patterns = {
        0: [r"\b(female|f|woman|girl)\b"],
        1: [r"\b(male|m|man|boy)\b"]
    }
    
    for value, patterns in gender_patterns.items():
        if any(re.search(pattern, text) for pattern in patterns):
            extracted_data["Gender"] = value
            break
    if "Gender" not in extracted_data:
        extracted_data["Gender"] = 0  # Default to female if not specified
    
    # Boolean fields with improved pattern matching
    boolean_fields = {
        "Smoking": [
            (r"\bsmoking[:\s]*(yes|y|positive|present|true|1)\b", 1),
            (r"\bsmoking[:\s]*(no|n|negative|absent|false|0)\b", 0)
        ],
        "Smoking History": [
            (r"\bsmoking\s*history[:\s]*(yes|y|positive|present|true|1)\b", 1),
            (r"\bsmoking\s*history[:\s]*(no|n|negative|absent|false|0)\b", 0)
        ],
        "Radiotherapy History": [
            (r"\bradiotherapy\s*history[:\s]*(yes|y|positive|present|true|1)\b", 1),
            (r"\bradiotherapy\s*history[:\s]*(no|n|negative|absent|none|false|0)\b", 0)
        ]
    }
    
    for field, patterns in boolean_fields.items():
        for pattern, value in patterns:
            if re.search(pattern, text):
                extracted_data[field] = value
                break
        if field not in extracted_data:
            extracted_data[field] = 0  # Default to 0 (No) if not found
    
    # Handle thyroid function field
    if re.search(r"\b(euthyroid|normal\s*thyroid)\b", text):
        extracted_data["Thyroid Function"] = 0
    elif re.search(r"\b(hypothyroid|hyperthyroid|abnormal\s*thyroid)\b", text):
        extracted_data["Thyroid Function"] = 1
    else:
        extracted_data["Thyroid Function"] = 0  # Default
    
    # Physical examination and adenopathy
    if re.search(r"\b(single\s*nodular\s*goiter|nodule\s*found)\b", text):
        extracted_data["Physical Examination"] = 1
    else:
        extracted_data["Physical Examination"] = 0
        
    if re.search(r"\b(adenopathy\s*present|enlarged\s*lymph\s*nodes)\b", text):
        extracted_data["Adenopathy"] = 1
    else:
        extracted_data["Adenopathy"] = 0
    
    # Cancer pathology type
    if re.search(r"\b(papillary\s*thyroid\s*carcinoma|ptc)\b", text):
        extracted_data["Types of Thyroid Cancer (Pathology)"] = 1
    elif re.search(r"\b(follicular)\b", text):
        extracted_data["Types of Thyroid Cancer (Pathology)"] = 2
    else:
        extracted_data["Types of Thyroid Cancer (Pathology)"] = 0
    
    # Focality
    if re.search(r"\b(multi-?focal|multiple\s*foci)\b", text):
        extracted_data["Focality"] = 1
    else:
        extracted_data["Focality"] = 0
    
    # Risk level
    if re.search(r"\brisk\s*category[:\s]*(high|3)\b", text):
        extracted_data["Risk"] = 2
    elif re.search(r"\brisk\s*category[:\s]*(intermediate|medium|2)\b", text):
        extracted_data["Risk"] = 1
    else:
        extracted_data["Risk"] = 0
    
    # TNM staging
    tnm_fields = {
        "Tumor": r"t\s*(\d+)",
        "Lymph Nodes": r"n\s*(\d+)",
        "Cancer Metastasis": r"m\s*(\d+)"
    }
    
    for field, pattern in tnm_fields.items():
        match = re.search(pattern, text)
        extracted_data[field] = int(match.group(1)) if match else 0
    
    # Stage (roman numerals)
    stage_match = re.search(r"stage\s*(i{1,3}|iv|v)\b", text)
    roman_to_int = {"i": 1, "ii": 2, "iii": 3, "iv": 4, "v": 5}
    if stage_match:
        extracted_data["Stage"] = roman_to_int.get(stage_match.group(1).lower(), 0)
    else:
        # Try numeric stage
        numeric_stage = re.search(r"stage\s*([1-5])\b", text)
        extracted_data["Stage"] = int(numeric_stage.group(1)) if numeric_stage else 0
    
    # Treatment response
    if re.search(r"\bresponse[:\s]*(structural\s*incomplete|poor|inadequate)\b", text):
        extracted_data["Treatment Response"] = 1
    elif re.search(r"\bresponse[:\s]*(excellent|good|complete)\b", text):
        extracted_data["Treatment Response"] = 0
    else:
        extracted_data["Treatment Response"] = 0
    
    print("🔍 Thyroid Extracted from report:", extracted_data)
    return extracted_data





# Create uploads directory for breast cancer module
bc_uploads = os.path.join(app.instance_path, 'bc_uploads')
os.makedirs(bc_uploads, exist_ok=True)

# Load the trained breast cancer model
bc_model = joblib.load('templates/breast_cancer_ensemble.pkl')

# Features required for breast cancer prediction
bc_features = ['radius_mean', 'texture_mean', 'perimeter_mean', 'area_mean', 'smoothness_mean']  # Example

# Function to extract text from PDF for breast cancer
def bc_extract_text_from_pdf(file):
    reader = PyPDF2.PdfReader(file)
    text = ''
    for page in reader.pages:
        text += page.extract_text() + ' '
    return text

# Function to extract text from DOCX for breast cancer
def bc_extract_text_from_docx(file):
    doc = docx.Document(file)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + ' '
    return text

# Function to extract data from text using regex for breast cancer
def bc_extract_data_from_text(text):
    extracted_data = {}
    patterns = {
    'radius_mean': r'Radius Mean:\s*([\d.]+)',
    'texture_mean': r'Texture Mean:\s*([\d.]+)',
    'perimeter_mean': r'Perimeter Mean:\s*([\d.]+)',
    'area_mean': r'Area Mean:\s*([\d.]+)',
    'smoothness_mean': r'Smoothness Mean:\s*([\d.]+)',
    'compactness_mean': r'Compactness Mean:\s*([\d.]+)',
    'concavity_mean': r'Concavity Mean:\s*([\d.]+)',
    'concave points_mean': r'Concave Points Mean:\s*([\d.]+)',
    'symmetry_mean': r'Symmetry Mean:\s*([\d.]+)',
    'fractal_dimension_mean': r'Fractal Dimension Mean:\s*([\d.]+)',
    'radius_se': r'Radius SE:\s*([\d.]+)',
    'texture_se': r'Texture SE:\s*([\d.]+)',
    'perimeter_se': r'Perimeter SE:\s*([\d.]+)',
    'area_se': r'Area SE:\s*([\d.]+)',
    'smoothness_se': r'Smoothness SE:\s*([\d.]+)',
    'compactness_se': r'Compactness SE:\s*([\d.]+)',
    'concavity_se': r'Concavity SE:\s*([\d.]+)',
    'concave points_se': r'Concave Points SE:\s*([\d.]+)',
    'symmetry_se': r'Symmetry SE:\s*([\d.]+)',
    'fractal_dimension_se': r'Fractal Dimension SE:\s*([\d.]+)',
    'radius_worst': r'Radius Worst:\s*([\d.]+)',
    'texture_worst': r'Texture Worst:\s*([\d.]+)',
    'perimeter_worst': r'Perimeter Worst:\s*([\d.]+)',
    'area_worst': r'Area Worst:\s*([\d.]+)',
    'smoothness_worst': r'Smoothness Worst:\s*([\d.]+)',
    'compactness_worst': r'Compactness Worst:\s*([\d.]+)',
    'concavity_worst': r'Concavity Worst:\s*([\d.]+)',
    'concave points_worst': r'Concave Points Worst:\s*([\d.]+)',
    'symmetry_worst': r'Symmetry Worst:\s*([\d.]+)',
    'fractal_dimension_worst': r'Fractal Dimension Worst:\s*([\d.]+)'
    }

    for feature, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        extracted_data[feature] = float(match.group(1)) if match else None

    return extracted_data

# Breast Cancer Routes
@app.route('/breast_cancer', methods=['GET'])
def breast_cancer_home():
    return render_template('breast_cancer.html', features=bc_features, extracted_data={}, prediction_text=None, precautions=None)

# File Upload Route for breast cancer - note the explicit route path
@app.route('/breast_cancer_extract', methods=['POST'])
def breast_cancer_extract():
    file = request.files['file']
    if file.filename.endswith('.pdf'):
        text = bc_extract_text_from_pdf(file)
    elif file.filename.endswith('.docx'):
        text = bc_extract_text_from_docx(file)
    else:
        return "Unsupported file format. Please upload a PDF or DOCX."

    extracted_data = bc_extract_data_from_text(text)
    return render_template('breast_cancer.html', features=bc_features, extracted_data=extracted_data, prediction_text=None, precautions=None)

# Prediction Route for breast cancer - note the explicit route path
@app.route('/breast_cancer_predict', methods=['POST'])
def breast_cancer_predict():
    input_data = {feature: request.form.get(feature) for feature in bc_features}

    # Convert input data for the model
    processed_data = []
    for feature in bc_features:
        value = input_data.get(feature)
        try:
            processed_data.append(float(value))
        except (TypeError, ValueError):
            processed_data.append(0.0)

    # Model prediction
    prediction = bc_model.predict([processed_data])[0]

    # Prediction result and precautions
    if prediction == 0:
        prediction_text = "Benign (No Cancer Detected)"
        precautions = ["Continue regular screenings and maintain a healthy lifestyle."]
    else:
        prediction_text = "Malignant (Cancer Detected)"
        precautions = [
            "Consult an oncologist for further tests.",
            "Follow recommended treatments and screenings.",
            "Maintain a healthy diet and exercise regularly."
        ]

    return render_template('breast_cancer.html', features=bc_features, extracted_data=input_data, prediction_text=prediction_text, precautions=precautions)






import pickle
import numpy as np

# Load lung cancer prediction model
with open("templates/lung_cancer_prediction_model.pkl", "rb") as model_file:
    lung_cancer_model = pickle.load(model_file)

# Column order for lung cancer model input
lung_cancer_columns = [
    "GENDER", "AGE", "SMOKING", "YELLOW_FINGERS", "ANXIETY",
    "PEER_PRESSURE", "CHRONIC DISEASE", "FATIGUE", "ALLERGY",
    "WHEEZING", "ALCOHOL CONSUMING", "COUGHING", "SHORTNESS OF BREATH",
    "SWALLOWING DIFFICULTY", "CHEST PAIN"
]

# Define the labels for the binary inputs (0/1)
# Using list instead of dict for Jinja2 compatibility
lung_cancer_binary_labels = {
    "GENDER": ["Female", "Male"],
    "SMOKING": ["No", "Yes"],
    "YELLOW_FINGERS": ["No", "Yes"],
    "ANXIETY": ["No", "Yes"],
    "PEER_PRESSURE": ["No", "Yes"],
    "CHRONIC DISEASE": ["No", "Yes"],
    "FATIGUE": ["No", "Yes"],
    "ALLERGY": ["No", "Yes"],
    "WHEEZING": ["No", "Yes"],
    "ALCOHOL CONSUMING": ["No", "Yes"],
    "COUGHING": ["No", "Yes"],
    "SHORTNESS OF BREATH": ["No", "Yes"],
    "SWALLOWING DIFFICULTY": ["No", "Yes"],
    "CHEST PAIN": ["No", "Yes"]
}

@app.route("/lung_cancer", methods=["GET", "POST"])
def lung_cancer():
    if request.method == "POST":
        try:
            # Get inputs and convert to model-compatible format
            user_inputs = [int(request.form[col]) for col in lung_cancer_columns]
            user_inputs = np.array(user_inputs).reshape(1, -1)

            # Make prediction
            prediction = lung_cancer_model.predict(user_inputs)[0]
            result = "Positive for Lung Cancer" if prediction == 1 else "Negative for Lung Cancer"
            
            # Prepare recommendations based on prediction
            if prediction == 1:
                recommendations = [
                    "Consult with a pulmonologist immediately for further evaluation",
                    "Complete diagnostic imaging tests as recommended by your physician",
                    "Follow up with appropriate specialist referrals",
                    "Consider joining a support group for emotional support"
                ]
            else:
                recommendations = [
                    "Continue regular health check-ups",
                    "Maintain a smoke-free lifestyle",
                    "Report any changes in respiratory symptoms to your doctor",
                    "Follow healthy lifestyle practices including regular exercise"
                ]

            return render_template("lung_cancer.html", 
                                  columns=lung_cancer_columns,
                                  binary_labels=lung_cancer_binary_labels,
                                  prediction=result,
                                  recommendations=recommendations,
                                  form_data=request.form)
        except Exception as e:
            return render_template("lung_cancer.html", 
                                  columns=lung_cancer_columns,
                                  binary_labels=lung_cancer_binary_labels,
                                  error=str(e))

    return render_template("lung_cancer.html", 
                          columns=lung_cancer_columns,
                          binary_labels=lung_cancer_binary_labels,
                          prediction=None)







import pandas as pd

# Liver Disease Model Loading
liver_model = joblib.load('templates/liver_disease.pkl')
print(type(liver_model))

# List of features expected from the form for liver disease
liver_features = [
    'age', 'gender', 'total_bilirubin', 'direct_bilirubin', 'alkaline_phosphotase',
    'alamine_aminotransferase', 'aspartate_aminotransferase', 'total_proteins',
    'albumin', 'albumin_and_globulin_ratio'
]

# Function to extract text from PDF (reuse if you already have this function)
def extract_text_from_pdf(file):
    reader = PyPDF2.PdfReader(file)
    text = ''
    for page in reader.pages:
        text += page.extract_text() + ' '
    return text

# Function to extract text from DOCX (reuse if you already have this function)
def extract_text_from_docx(file):
    doc = docx.Document(file)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + ' '
    return text

# Function to extract liver data from text using regex
def extract_liver_data_from_text(text):
    extracted_data = {}
    patterns = {
        'age': r'Age:\s*(\d+)',
        'gender': r'Gender:\s*(Male|Female)',
        'total_bilirubin': r'Total Bilirubin:\s*([\d.]+)',
        'direct_bilirubin': r'Direct Bilirubin:\s*([\d.]+)',
        'alkaline_phosphotase': r'Alkaline Phosphatase:\s*(\d+)',
        'alamine_aminotransferase': r'ALT:\s*(\d+)',
        'aspartate_aminotransferase': r'AST:\s*(\d+)',
        'total_proteins': r'Total Proteins:\s*([\d.]+)',
        'albumin': r'Albumin:\s*([\d.]+)',
        'albumin_and_globulin_ratio': r'Albumin & Globulin Ratio:\s*([\d.]+)'
    }
    
    for feature, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        extracted_data[feature] = match.group(1) if match else None
    
    return extracted_data

# Liver Disease Home Route
@app.route('/liver', methods=['GET'])
def liver_home():
    return render_template('liver_disease.html', features=liver_features, extracted_data={}, prediction_text=None, precautions=None)

# Liver Disease File Upload Route
@app.route('/extract/liver', methods=['POST'])
def extract_liver():
    file = request.files['file']
    if file.filename.endswith('.pdf'):
        text = extract_text_from_pdf(file)
    elif file.filename.endswith('.docx'):
        text = extract_text_from_docx(file)
    else:
        return "Unsupported file format. Please upload a PDF or DOCX."
    
    extracted_data = extract_liver_data_from_text(text)
    return render_template('liver_disease.html', features=liver_features, extracted_data=extracted_data, prediction_text=None, precautions=None)

# Liver Disease Prediction Route
@app.route('/predict/liver', methods=['POST'])
def predict_liver():
    input_data = {feature: request.form.get(feature) for feature in liver_features}
    
    # Convert input data for the model
    processed_data = []
    for feature in liver_features:
        value = input_data.get(feature)
        if value and value.lower() in ['male']:
            processed_data.append(1)  # Encode Male as 1
        elif value and value.lower() in ['female']:
            processed_data.append(0)  # Encode Female as 0
        else:
            try:
                processed_data.append(float(value))  # Convert numbers to float
            except (TypeError, ValueError):
                processed_data.append(0.0)  # Default to 0.0 if invalid input
    
    # Model prediction
    input_array = np.array(processed_data).reshape(1, -1)
    prediction = liver_model.predict(input_array)[0]
    
    # Interpret Prediction
    prediction_text = "Positive (Liver Disease Detected)" if prediction == 1 else "Negative (No Liver Disease)"
    precautions = [
        "Limit alcohol intake and eat a balanced diet.",
        "Exercise regularly and maintain a healthy weight.",
        "Monitor liver enzyme levels periodically."
    ] if prediction == 1 else ["No immediate precautions needed. Continue regular health check-ups."]
    
    return render_template('liver_disease.html', features=liver_features, extracted_data=input_data, prediction_text=prediction_text, precautions=precautions)

if __name__ == '__main__':
    if not os.path.exists('uploads'):
        os.makedirs('uploads')
    app.run(debug=True)
